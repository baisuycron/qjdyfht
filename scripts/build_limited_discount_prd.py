from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parent.parent / "限时折扣需求文档.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
GRAY = "666666"
WHITE = "FFFFFF"
CONTENT_WIDTH_DXA = 9360


def set_run_font(run, size=11, bold=False, color="000000", name="Microsoft YaHei"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths, indent=120):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for index, width in enumerate(widths):
        grid.gridCol_lst[index].set(qn("w:w"), str(width))
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table, color="D9E2F0"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        edge_el = borders.find(tag)
        if edge_el is None:
            edge_el = OxmlElement(f"w:{edge}")
            borders.append(edge_el)
        edge_el.set(qn("w:val"), "single")
        edge_el.set(qn("w:sz"), "4")
        edge_el.set(qn("w:space"), "0")
        edge_el.set(qn("w:color"), color)


def write_cell(cell, text, bold=False, color="000000", size=10.5, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, color=color)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table)
    header = table.rows[0]
    header.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    for index, heading in enumerate(headers):
        set_cell_shading(header.cells[index], LIGHT_BLUE)
        write_cell(header.cells[index], heading, bold=True, color=INK, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            write_cell(cells[index], value, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER if index == 0 else WD_ALIGN_PARAGRAPH.LEFT)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    return table


def add_body(doc, text, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    r = p.add_run(text)
    set_run_font(r, size=11, color="000000")
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    r = p.add_run(text)
    set_run_font(r, size=11)
    return p


def add_h1(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run_font(r, size=16, bold=True, color=BLUE)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run_font(r, size=13, bold=True, color=BLUE)
    return p


def add_callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    set_table_borders(table, color="D5E1F0")
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(f"{title}：")
    set_run_font(r, size=10.5, bold=True, color=DARK_BLUE)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color="333333")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_meta_row(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.1
    left = p.add_run(f"{label}：")
    set_run_font(left, size=10.5, bold=True, color=INK)
    right = p.add_run(value)
    set_run_font(right, size=10.5, color="333333")


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    for style_name, size, color in (("Heading 1", 16, BLUE), ("Heading 2", 13, BLUE), ("Heading 3", 12, DARK_BLUE)):
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)

    # memo_masthead: restrained running label and title metadata.
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run("千金大药房 | 产品需求文档")
    set_run_font(run, size=9, color=GRAY)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(0)
    run = footer.add_run("第 ")
    set_run_font(run, size=9, color=GRAY)
    add_page_field(footer)
    run = footer.add_run(" 页")
    set_run_font(run, size=9, color=GRAY)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(10)
    title.paragraph_format.space_after = Pt(4)
    title.paragraph_format.keep_with_next = True
    r = title.add_run("限时折扣需求文档")
    set_run_font(r, size=23, bold=True, color=INK)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(14)
    r = subtitle.add_run("适用于管理后台展示与小程序端促销呈现")
    set_run_font(r, size=12, color=GRAY)
    add_meta_row(doc, "所属模块", "营销中心 > 门店促销 > 限时折扣")
    add_meta_row(doc, "版本", "V1.0")
    add_meta_row(doc, "日期", "2026-08-05")
    add_meta_row(doc, "维护边界", "活动规则、商品、门店及适用对象均由 ERP 维护")
    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(8)
    rule.paragraph_format.space_after = Pt(0)
    ppr = rule._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), BLUE)
    borders.append(bottom)
    ppr.append(borders)

    add_h1(doc, "一、需求背景")
    add_body(doc, "当前门店促销活动由 ERP 统一维护。为保证小程序端的促销信息、结算价格与 ERP 配置一致，管理后台需提供限时折扣活动的查询与详情查看能力，并在小程序商品侧以统一、易理解的形式展示优惠。")
    add_body(doc, "本需求不新增独立的活动配置能力。活动的创建、编辑、删除、门店与商品范围、适用对象、权限审批及数据维护均以 ERP 为唯一来源。")
    add_callout(doc, "核心原则", "管理后台负责展示与查询，ERP 负责维护；小程序端和结算端均以 ERP 同步的有效活动数据为准。")

    add_h1(doc, "二、需求目标")
    for text in (
        "让运营与门店人员能够快速查询限时折扣活动的生效范围、时间、状态及参与商品。",
        "让消费者在小程序端清楚识别限时优惠类型，减少价格理解成本。",
        "保障促销价与会员价并存时，订单按更低价格自动生效。",
        "明确系统边界，避免在管理后台重复维护 ERP 已维护的活动规则。",
    ):
        add_bullet(doc, text)

    add_h1(doc, "三、业务流程")
    add_body(doc, "限时折扣活动从 ERP 维护开始，经数据同步后分别支撑管理后台查询和小程序端展示、结算。")
    add_table(
        doc,
        ["步骤", "处理方", "业务动作", "产出"],
        [
            ("1", "ERP", "维护活动、优惠方式、商品、门店、时间与适用对象。", "可同步的限时折扣活动数据"),
            ("2", "数据同步", "将活动及参与商品明细同步至营销服务。", "有效活动与商品优惠规则"),
            ("3", "管理后台", "按名称、编号、活动时间查询活动，并查看详情。", "活动清单与明细"),
            ("4", "小程序", "商品命中有效活动时展示对应限时优惠标识。", "限时打折/特价/减价提示"),
            ("5", "结算", "比较活动优惠价与会员价，取更低价格。", "实际成交价"),
        ],
        [700, 1300, 4200, 3160],
    )
    add_callout(doc, "冲突处理", "同一商品在同一时间只能参加一个促销活动；若同时存在会员价，系统比较会员价与该促销活动计算后的价格，按更低价格生效。")

    add_h1(doc, "四、产品方案概要")
    add_table(
        doc,
        ["范围", "方案说明"],
        [
            ("管理后台", "提供限时折扣活动列表、条件查询、重置和活动详情查看，不提供新增、编辑、删除、停用等维护操作。"),
            ("小程序端", "在商品列表、商品详情及购物车等承载商品价格的页面，展示与优惠方式匹配的限时优惠文案。"),
            ("结算规则", "商品仅命中一个促销活动；活动优惠价与会员价同时存在时，自动取更低价格。"),
            ("ERP 边界", "活动、商品、门店、时间、参与日期、参与星期、适用对象、权限与审批均由 ERP 维护并同步。"),
        ],
        [2500, 6860],
    )

    add_h1(doc, "五、产品方案细节")
    add_h2(doc, "5.1 管理后台 - 活动列表")
    add_body(doc, "入口路径：营销中心 > 门店促销 > 限时折扣。页面默认展示 ERP 已同步的限时折扣活动。")
    add_table(
        doc,
        ["区域", "字段/动作", "规则"],
        [
            ("查询条件", "活动名称", "支持模糊查询。"),
            ("查询条件", "活动编号", "支持精确或包含查询，以 ERP 活动编号为准。"),
            ("查询条件", "活动时间", "按活动起止时间筛选，与筛选时间区间有交集的活动应被返回。"),
            ("列表字段", "名称、编号、促销类型、活动时间、适用门店、活动状态、最后修改时间", "字段来自 ERP 同步数据；适用门店显示“全部门店”或“指定门店”。"),
            ("操作", "查看、重置", "查看进入活动详情；重置清空全部查询条件并恢复默认列表。"),
        ],
        [1400, 3300, 4660],
    )

    add_h2(doc, "5.2 管理后台 - 活动详情")
    add_body(doc, "详情页为只读页面，用于核对 ERP 已配置的活动规则与参与商品。")
    add_table(
        doc,
        ["信息分组", "展示字段"],
        [
            ("基本信息", "活动名称、活动编号、优惠方式、活动状态、活动时间、适用门店、参与日期、参与星期、最后修改时间。"),
            ("参与商品明细", "商品编码、优惠方式、优惠内容；展示活动内全部已同步商品。"),
            ("优惠内容", "按折扣率展示 X 折；按促销价展示促销价；按减价格展示立减金额。"),
        ],
        [2500, 6860],
    )

    add_h2(doc, "5.3 小程序端 - 展示规则")
    add_table(
        doc,
        ["优惠方式", "小程序展示文案", "展示参数"],
        [
            ("按折扣率", "限时打 X 折", "X 为 ERP 下发折扣率换算后的折数，例如 0.88 显示为 8.8 折。"),
            ("按促销价", "限时特价", "同时展示活动计算后的促销价格。"),
            ("按减价格", "限时减 X", "X 为 ERP 下发的减免金额。"),
        ],
        [1800, 2500, 5060],
    )
    add_body(doc, "仅当商品在当前门店、当前时间和当前适用条件下命中有效活动时展示限时优惠。未命中或活动结束后，不展示对应标识。")

    add_h2(doc, "5.4 价格生效规则")
    add_table(
        doc,
        ["规则项", "规则说明"],
        [
            ("活动唯一性", "同一商品在同一时间只能参加一个促销活动。活动冲突由 ERP 维护侧控制。"),
            ("活动价计算", "按折扣率：销售价 × 折扣率；按促销价：使用 ERP 下发促销价；按减价格：销售价 - 减免金额。"),
            ("会员价比较", "当商品同时具备活动优惠价与会员价时，比较两者，按更低价格作为实际成交价。"),
            ("展示与结算一致", "商品展示价格、购物车价格和提交订单价格须使用同一套价格计算结果。"),
        ],
        [2500, 6860],
    )

    add_h2(doc, "5.5 状态与边界")
    for text in (
        "活动状态按活动起止时间计算并展示为：未开始、进行中、已结束。",
        "ERP 配置变更后的同步时效、失败重试、字段映射与告警策略由技术方案另行确认；在同步完成前，前台继续使用最近一次成功同步的数据。",
        "本期不包含后台活动创建、编辑、复制、删除、提前结束、审批、权限配置及 ERP 配置回写。",
    ):
        add_bullet(doc, text)

    add_h2(doc, "5.6 验收标准")
    add_table(
        doc,
        ["编号", "验收项", "通过标准"],
        [
            ("AC-01", "活动查询", "可按活动名称、活动编号和活动时间查询；点击重置后恢复默认列表。"),
            ("AC-02", "活动详情", "详情完整展示活动基础规则与参与商品明细，且不可编辑。"),
            ("AC-03", "优惠文案", "三种优惠方式分别展示“限时打 X 折”“限时特价”“限时减 X”。"),
            ("AC-04", "促销冲突", "同一商品同一时间仅命中一个促销活动，冲突配置由 ERP 侧避免。"),
            ("AC-05", "会员价", "活动优惠价与会员价并存时，商品展示和结算均按更低价格生效。"),
            ("AC-06", "边界控制", "管理后台不提供活动维护入口，ERP 为活动数据唯一维护来源。"),
        ],
        [1100, 2500, 5760],
    )

    doc.core_properties.title = "限时折扣需求文档"
    doc.core_properties.subject = "千金大药房限时折扣需求"
    doc.core_properties.author = "千金大药房"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_document()
