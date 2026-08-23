from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from build_combination_purchase_prds import GRAY, add_page_field, set_run_font


ROOT = Path(__file__).resolve().parent.parent
ADMIN_PATH = ROOT / "后台系统组合购PRD_V1.2.docx"
MALL_PATH = ROOT / "商城端组合购PRD_V1.2.docx"
OUTPUT_PATH = ROOT / "组合购PRD_后台系统及商城端_V1.2.docx"


def clear_container(container):
    paragraphs = list(container.paragraphs)
    first = paragraphs[0]
    first.clear()
    for paragraph in paragraphs[1:]:
        paragraph._element.getparent().remove(paragraph._element)
    return first


def configure_section_chrome(section, short_title):
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False

    header = clear_container(section.header)
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    set_run_font(header.add_run("千金健康商城 | 产品需求文档"), size=9, color=GRAY)

    footer = clear_container(section.footer)
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_after = Pt(0)
    set_run_font(footer.add_run(f"{short_title}  ·  第 "), size=9, color=GRAY)
    add_page_field(footer)
    set_run_font(footer.add_run(" 页"), size=9, color=GRAY)


def append_body_before_final_section(target, source):
    target_body = target.element.body
    final_section_properties = target_body.sectPr
    for element in source.element.body:
        if element.tag.endswith("}sectPr"):
            continue
        final_section_properties.addprevious(deepcopy(element))


def validate_merged_document(path, admin, mall):
    merged = Document(path)
    merged_paragraphs = [paragraph.text for paragraph in merged.paragraphs]
    admin_paragraphs = [paragraph.text for paragraph in admin.paragraphs]
    mall_paragraphs = [paragraph.text for paragraph in mall.paragraphs]
    paragraphs = [text.strip() for text in merged_paragraphs]
    admin_title = "千金健康商城后台系统｜组合购 PRD"
    mall_title = "千金健康商城端｜组合购 PRD"

    assert admin_title in paragraphs, "缺少后台系统组合购标题"
    assert mall_title in paragraphs, "缺少商城端组合购标题"
    assert paragraphs.index(admin_title) < paragraphs.index(mall_title), "文档顺序错误"
    assert len(merged.sections) == 2, f"分节数量异常：{len(merged.sections)}"
    assert len(merged.tables) == len(admin.tables) + len(mall.tables), "表格数量不完整"
    assert len(merged.paragraphs) == len(admin.paragraphs) + len(mall.paragraphs) + 1, "段落数量不完整"
    assert merged_paragraphs[: len(admin_paragraphs)] == admin_paragraphs, "后台正文内容或顺序发生变化"
    assert merged_paragraphs[len(admin_paragraphs)] == "", "两部分之间缺少独立分节段落"
    assert merged_paragraphs[len(admin_paragraphs) + 1 :] == mall_paragraphs, "商城正文内容或顺序发生变化"

    table_signature = lambda table: [[cell.text for cell in row.cells] for row in table.rows]
    merged_tables = [table_signature(table) for table in merged.tables]
    admin_tables = [table_signature(table) for table in admin.tables]
    mall_tables = [table_signature(table) for table in mall.tables]
    assert merged_tables == admin_tables + mall_tables, "表格内容或顺序发生变化"

    first_footer = merged.sections[0].footer.paragraphs[0].text
    second_footer = merged.sections[1].footer.paragraphs[0].text
    assert "后台组合购 PRD" in first_footer, "后台部分页脚错误"
    assert "商城组合购 PRD" in second_footer, "商城部分页脚错误"

    return {
        "paragraphs": len(merged.paragraphs),
        "tables": len(merged.tables),
        "sections": len(merged.sections),
        "admin_title_index": paragraphs.index(admin_title),
        "mall_title_index": paragraphs.index(mall_title),
    }


def main():
    admin = Document(ADMIN_PATH)
    mall = Document(MALL_PATH)

    merged = Document(ADMIN_PATH)
    configure_section_chrome(merged.sections[0], "后台组合购 PRD")
    mall_section = merged.add_section(WD_SECTION.NEW_PAGE)
    configure_section_chrome(mall_section, "商城组合购 PRD")
    append_body_before_final_section(merged, mall)

    merged.core_properties.title = "组合购 PRD｜后台系统及商城端"
    merged.core_properties.subject = "后台系统组合购与商城端组合购完整产品需求"
    merged.core_properties.author = "千金健康商城产品团队"
    merged.core_properties.keywords = "组合购, 后台系统, 商城端, 固定组合价, PRD"
    merged.save(OUTPUT_PATH)

    summary = validate_merged_document(OUTPUT_PATH, admin, mall)
    print(OUTPUT_PATH)
    print(summary)


if __name__ == "__main__":
    main()
