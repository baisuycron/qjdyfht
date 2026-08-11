from pathlib import Path
import sys

from docx import Document


path = Path(sys.argv[1])
doc = Document(path)
headings = [p.text.strip() for p in doc.paragraphs if p.style.name.startswith("Heading")]
required = [
    "一、需求背景",
    "二、需求目标",
    "三、业务流程",
    "四、产品方案概要",
    "五、产品方案细节",
]
text = "\n".join(p.text for p in doc.paragraphs)
text += "\n" + "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)
errors = []
if [heading for heading in headings if heading in required] != required:
    errors.append("five top-level headings are missing or out of order")
if any(len(table.rows) == 0 or len(table.columns) == 0 for table in doc.tables):
    errors.append("empty table detected")
if "\ufffd" in text:
    errors.append("replacement character detected")
for phrase in (
    "warecode", "warename", "warespec", "wareunit", "factoryname", "marketingType",
    "cnt", "itemList.wareid", "itemList.wareqty", "queryMallCollocationList",
    "collocationId", "wareqty × goodCount", "AC-10",
):
    if phrase not in text:
        errors.append(f"required content missing: {phrase}")
print(f"paragraphs={len(doc.paragraphs)}")
print(f"tables={len(doc.tables)}")
print(f"headings={len(headings)}")
print("top_level=" + " | ".join(heading for heading in headings if heading in required))
print(f"replacement_chars={text.count(chr(0xFFFD))}")
print("status=" + ("FAIL: " + "; ".join(errors) if errors else "OK"))
raise SystemExit(1 if errors else 0)
