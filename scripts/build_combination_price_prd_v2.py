from pathlib import Path
import sys

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(sys.argv[1]) if len(sys.argv) > 1 else Path(__file__).resolve().parent.parent / "组合价PRD.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
GRAY = "666666"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(run, size=11, bold=False, color="000000", italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def configure_style(style, size, color="000000", bold=False, before=0, after=6, line=1.10):
    style.font.name = "Calibri"
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = line


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_geometry(table, widths, indent=TABLE_INDENT_DXA):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    for tag, value in (("w:tblW", sum(widths)), ("w:tblInd", indent)):
        node = tbl_pr.first_child_found_in(tag)
        if node is None:
            node = OxmlElement(tag)
            tbl_pr.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    for index, width in enumerate(widths):
        table._tbl.tblGrid.gridCol_lst[index].set(qn("w:w"), str(width))
    for row in table.rows:
        prevent_row_split(row)
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table, color="D9DEE7", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def write_cell(cell, text, bold=False, color="000000", size=10.2, align=WD_ALIGN_PARAGRAPH.LEFT):
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.10
    run = paragraph.add_run(str(text))
    set_run_font(run, size=size, bold=bold, color=color)


def add_table(doc, headers, rows, widths, center_columns=None):
    center_columns = set(center_columns or [])
    table = doc.add_table(rows=1, cols=len(headers))
    for index, label in enumerate(headers):
        set_cell_shading(table.rows[0].cells[index], LIGHT_GRAY)
        write_cell(table.rows[0].cells[index], label, bold=True, color=INK, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_repeat_table_header(table.rows[0])
    for row_values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            alignment = WD_ALIGN_PARAGRAPH.CENTER if index in center_columns else WD_ALIGN_PARAGRAPH.LEFT
            write_cell(cells[index], value, align=alignment)
    set_table_geometry(table, widths)
    set_table_borders(table)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(2)
    return table


def add_body(doc, text, bold_lead=None):
    paragraph = doc.add_paragraph(style="Normal")
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_run_font(lead, bold=True, color=INK)
        body = paragraph.add_run(text[len(bold_lead):])
        set_run_font(body)
    else:
        set_run_font(paragraph.add_run(text))
    return paragraph


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    if level == 1:
        set_run_font(run, size=16, bold=True, color=BLUE)
    elif level == 2:
        set_run_font(run, size=13, bold=True, color=BLUE)
    else:
        set_run_font(run, size=12, bold=True, color=DARK_BLUE)
    return paragraph


def next_num_id(numbering):
    ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    return max(ids, default=0) + 1


def next_abstract_num_id(numbering):
    ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    return max(ids, default=0) + 1


def create_list_numbering(doc, ordered=False):
    numbering = doc.part.numbering_part.element
    abstract_id = next_abstract_num_id(numbering)
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if ordered else "bullet")
    level.append(num_fmt)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1." if ordered else "•")
    level.append(level_text)
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    level.append(justification)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "720")
    indent.set(qn("w:hanging"), "360")
    p_pr.append(indent)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)
    num_id = next_num_id(numbering)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_num_id = OxmlElement("w:abstractNumId")
    abstract_num_id.set(qn("w:val"), str(abstract_id))
    num.append(abstract_num_id)
    numbering.append(num)
    return num_id


def add_list_item(doc, text, num_id):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.167
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.extend((ilvl, num_id_node))
    p_pr.append(num_pr)
    set_run_font(paragraph.add_run(text))
    return paragraph


def add_callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    set_repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.10
    set_run_font(paragraph.add_run(f"{title}："), size=10.5, bold=True, color=DARK_BLUE)
    set_run_font(paragraph.add_run(text), size=10.5, color="333333")
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    set_table_borders(table, color="D5DCE6")
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_meta_row(doc, label, value):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.10
    set_run_font(paragraph.add_run(f"{label}："), size=10.5, bold=True, color=INK)
    set_run_font(paragraph.add_run(value), size=10.5, color="333333")


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, text, end))


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    configure_style(doc.styles["Normal"], 11, after=6, line=1.10)
    configure_style(doc.styles["Heading 1"], 16, BLUE, True, before=16, after=8, line=1.0)
    configure_style(doc.styles["Heading 2"], 13, BLUE, True, before=12, after=6, line=1.0)
    configure_style(doc.styles["Heading 3"], 12, DARK_BLUE, True, before=8, after=4, line=1.0)
    bullet_num_id = create_list_numbering(doc, ordered=False)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    set_run_font(header.add_run("千金大药房 | 产品需求文档"), size=9, color=GRAY)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_after = Pt(0)
    set_run_font(footer.add_run("第 "), size=9, color=GRAY)
    add_page_field(footer)
    set_run_font(footer.add_run(" 页"), size=9, color=GRAY)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(12)
    title.paragraph_format.space_after = Pt(4)
    title.paragraph_format.keep_with_next = True
    set_run_font(title.add_run("千金小程序需求-组合价"), size=23, bold=True, color=INK)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(14)
    subtitle.paragraph_format.keep_with_next = True
    set_run_font(subtitle.add_run("ERP维护、管理后台查询与小程序组合购"), size=12, color=GRAY)
    add_meta_row(doc, "所属模块", "营销中心 > 门店促销 > 组合价")
    add_meta_row(doc, "文档版本", "V1.1")
    add_meta_row(doc, "文档日期", "2026-08-11")
    add_meta_row(doc, "维护边界", "组合规则由ERP维护；管理后台负责只读查询；小程序负责组合选择、订单预览与交易展示")

    add_heading(doc, "一、需求背景", 1)
    add_body(doc, "当前组合商品由ERP维护。管理后台需要通过ERP开放接口查询组合基本资料和组成商品明细；小程序商城已有“组合购”代码链路，可按商品查询可用组合、选择组合商品与规格，并直接进入组合订单预览。")
    add_body(doc, "本需求以商城项目 wechat-miniapp 中现有组合购实现为依据，同时保留接口文档明确的字段边界。ERP分页查询接口主要服务管理后台核对；商城交易使用独立的组合购接口与订单预览接口，两类接口不能混用，也不能用Mock字段补齐缺失数据。")
    add_callout(doc, "现状结论", "商城已实现组合选择、价格预估、规格与库存校验、组合订单预览和提交，但当前编译版商品详情未发现触发组合购弹层的可见入口；管理后台接口也不返回组合成交价、活动时间、门店范围或交易状态。上述缺口需在本需求中明确补齐或确认。")

    add_heading(doc, "二、需求目标", 1)
    goals = (
        "让运营人员在管理后台按ERP接口支持范围查询组合资料和组成商品，避免产生第二套维护入口。",
        "在小程序商品详情提供明确的组合价入口，沿用现有组合购弹层、套餐选择、规格选择和数量校验逻辑。",
        "让消费者清楚看到组合商品、套餐价、原价、最高可省金额和当前选择的优惠金额。",
        "组合购买直接进入专用订单预览，订单金额以服务端预览结果为准，防止前端展示价与成交价不一致。",
        "识别ERP每组数量与商城统一套数之间的映射缺口，保证每个商品的实际购买数量计算正确。",
        "明确组合价与购物车、优惠券、会员价、配送方式、退款和换货的边界及待确认事项。",
    )
    for goal in goals:
        add_list_item(doc, goal, bullet_num_id)

    add_heading(doc, "三、业务流程", 1)
    add_body(doc, "组合价从ERP维护开始，分别支撑管理后台核对与小程序交易。小程序现有实现为“商品详情进入组合购，选好商品后直接进入订单预览”，不经过普通购物车。")
    add_table(doc, ["步骤", "处理方", "业务动作", "结果"], [
        ("1", "ERP", "维护组合主数据、组成商品及每组数量；组合规则仍以ERP为唯一维护来源。", "形成组合主数据与绑定明细。"),
        ("2", "管理后台接口", "按绑定明细最后修改时间分页查询组合主表，并为每条组合返回完整itemList。", "供管理后台列表和详情核对。"),
        ("3", "商品详情", "查询当前商品可参加的组合列表；有可用组合时展示组合价入口，默认进入第一组。", "打开组合购选择弹层。"),
        ("4", "组合选择", "主商品必选；消费者可勾选其他商品、选择规格和组合套数，页面计算套餐价与优惠金额。", "形成选中SKU清单。"),
        ("5", "购买校验", "校验至少选择2个商品、每个商品已选规格且库存满足购买数量。", "校验通过后生成collocationId与skuList。"),
        ("6", "订单预览", "调用组合购专用预览接口，结合配送方式、地址/自提门店重新计算商品、运费、优惠和应付金额。", "返回权威订单预览结果。"),
        ("7", "订单提交", "订单页保持组合ID和固定商品数量，使用预览信息与submitToken提交订单。", "生成可追溯的组合订单。"),
    ], [700, 1300, 4500, 2860], center_columns={0})
    add_callout(doc, "关键口径", "组合弹层中的金额只用于即时展示；订单成交金额、优惠资格、配送可用性和库存均以组合订单预览接口的返回为准。")

    add_heading(doc, "四、产品方案概要", 1)
    add_table(doc, ["范围", "产品方案"], [
        ("ERP", "维护组合资料、组成商品及每组数量；不在管理后台或小程序重复提供维护能力。"),
        ("管理后台", "按ERP分页查询接口提供列表、查询、重置、分页和只读详情，仅展示接口真实返回字段。"),
        ("小程序商品详情", "有可用组合时展示组合价入口；打开后加载套餐列表，默认选择第一组和主商品。"),
        ("组合选择弹层", "支持套餐切换、可选商品勾选、规格选择和统一套数调整；展示最高省、套餐价和当前优惠金额。"),
        ("订单预览与提交", "选中商品直接进入组合订单预览，不加入普通购物车；预览成功后沿用统一订单页完成配送选择和提交。"),
        ("订单/售后", "订单需保留组合活动标识和商品明细；组合专用退款、换货及金额分摊规则尚需业务确认。"),
    ], [2100, 7260])
    add_body(doc, "产品名称统一使用“组合价”，消费者操作文案可保留商城现有“组合购”“套餐价”“组合购买”。管理后台不展示“ERP同步·只读”“接口字段原型”“Mock数据”等徽标；只读属性通过不提供维护按钮体现。")

    doc.add_page_break()
    add_heading(doc, "五、产品方案细节", 1)

    add_heading(doc, "5.1 管理后台-功能入口与权限", 2)
    add_body(doc, "入口：营销中心 > 门店促销 > 组合价。进入后默认展示组合价列表，点击“查看”进入只读详情，点击“返回列表”返回原查询结果。")
    add_table(doc, ["项目", "要求"], [
        ("可用操作", "查询、重置、分页、查看、返回列表。"),
        ("禁止操作", "不提供新增、编辑、复制、删除、启用、停用、审批或回写ERP。"),
        ("页面提示", "不展示ERP同步、只读、接口原型或Mock数据徽标。"),
        ("数据来源", "生产环境以ERP组合价分页查询接口返回为准；前端不得自行补全接口未返回字段。"),
    ], [2200, 7160])

    add_heading(doc, "5.2 管理后台-查询条件与接口请求", 2)
    add_body(doc, "接口地址：https://wx.qjdyf.cn/external-dataservice/marketing-activities/combination-price/query。调用方式为POST，由管理后台服务端完成签名并请求接口，SECRET_KEY不得下发到浏览器。")
    add_table(doc, ["参数", "位置", "是否必填", "产品规则"], [
        ("appId", "Header", "是", "调用方应用ID，由服务端配置。"),
        ("sign", "Header", "是", "服务端按MD5(appId + SECRET_KEY + timestamp)生成。"),
        ("timestamp", "Header", "是", "秒级时间戳。"),
        ("startTime", "Body", "是", "查询开始时间，包含该时间点。页面选择开始日期后传当日00:00:00。"),
        ("endTime", "Body", "是", "查询结束时间，不包含该时间点。页面按自然日查询时，所选结束日期转换为次日00:00:00。"),
        ("pageNo", "Body", "是", "页码从1开始；查询条件或每页数量变化时回到第1页。"),
        ("pageSize", "Body", "是", "每页组合主表数量，接口最大500；页面提供10、20、50、100条/页。"),
    ], [1600, 1300, 1300, 5160], center_columns={1, 2})
    add_body(doc, "查询按钮：同时提交时间范围、pageNo和pageSize。重置按钮：清空页面时间条件并恢复默认页码和每页数量；若生产接口要求时间必填，重置后的首次查询应由服务端补充默认增量时间窗或在页面提示选择完整范围。")

    add_heading(doc, "5.3 管理后台-组合价列表", 2)
    add_table(doc, ["显示项", "接口字段", "展示规则"], [
        ("组合编码", "warecode", "按字符串展示，保留前导0；为空显示“—”。"),
        ("组合名称", "warename", "完整展示或在列宽内省略，悬停可查看完整值；为空显示“—”。"),
        ("规格/单位", "warespec / wareunit", "两者都有时显示“规格 / 单位”；只有一个时展示已有值。"),
        ("厂家", "factoryname", "为空显示“—”。"),
        ("营销类型", "marketingType", "按接口原值展示，当前示例为“组合价”。"),
        ("组成商品数", "cnt", "展示接口返回的组合商品数量；为空时可使用itemList实际条数作为兜底并记录异常。"),
        ("操作", "—", "仅提供“查看”。"),
    ], [1900, 2100, 5360])
    add_body(doc, "列表不展示：最后修改时间、组合售价、活动状态、活动时间、适用门店、渠道、库存、图片等接口未返回字段。")

    add_heading(doc, "5.4 管理后台-组合价详情", 2)
    add_body(doc, "详情页分为“基本信息”和“组合商品明细”两部分，不提供任何维护按钮或状态徽标。")
    add_table(doc, ["区域", "显示项", "接口字段", "展示规则"], [
        ("基本信息", "组合名称", "warename", "为空显示“—”。"),
        ("基本信息", "组合编码", "warecode", "按字符串展示并保留前导0。"),
        ("基本信息", "组合ID", "wareid", "展示组合主表ID。"),
        ("基本信息", "营销类型", "marketingType", "按接口原值展示。"),
        ("基本信息", "组合规格", "warespec", "为空显示“—”。"),
        ("基本信息", "组合单位", "wareunit", "为空显示“—”。"),
        ("基本信息", "组合厂家", "factoryname", "为空显示“—”。"),
        ("基本信息", "组成商品数", "cnt", "展示接口返回数量。"),
        ("商品明细", "商品ID", "itemList.wareid", "展示绑定商品ID。"),
        ("商品明细", "商品名称", "itemList.warename", "为空显示“—”。"),
        ("商品明细", "每组数量", "itemList.wareqty", "展示每一组组合所需数量，不解释为库存或限购数量。"),
    ], [1500, 1900, 2300, 3660])
    add_body(doc, "商品明细不展示商品编码和商品规格，因为当前接口未返回这两个字段。itemList按接口返回顺序展示，不做独立分页。")

    add_heading(doc, "5.5 小程序端-组合价功能入口", 2)
    add_body(doc, "商城现有组合购组件已挂载在商品详情页，并提供handleOpenCombinationPurchase方法，但当前编译模板未发现调用该方法的可见按钮或活动行。因此本期必须补齐真实入口，不能仅保留不可触达的弹层组件。")
    add_table(doc, ["场景", "产品规则"], [
        ("入口显示", "商品详情加载后，根据当前商品、门店和用户上下文查询可用组合。存在至少1个可交易组合时展示“组合价”优惠行或“组合购”按钮；无组合时隐藏。"),
        ("入口点击", "点击后打开现有组合购底部弹层，并以当前商品productId调用queryMallCollocationList。"),
        ("加载中", "入口点击后展示加载状态；不得先展示虚构套餐或价格。"),
        ("无可用组合", "接口返回空列表时关闭弹层并提示“当前商品暂无可用组合”。"),
        ("接口失败", "保留商品详情页正常购买能力，提示“组合价加载失败，请稍后重试”。"),
    ], [2200, 7160])
    add_callout(doc, "接口分工", "小程序使用商城接口 histore-gw/mallApi/apiCollocation/queryMallCollocationList 查询可交易组合；管理后台使用ERP开放分页接口查询主数据。商城不得直接使用后台分页结果进行计价。")

    add_heading(doc, "5.6 小程序端-组合商品选择", 2)
    add_table(doc, ["区域/动作", "商城现有逻辑", "本期产品要求"], [
        ("套餐切换", "展示combinationList中的套餐标题，默认选择第一组。", "切换套餐后刷新商品、价格、SKU选择和组合ID；不同套餐的选择状态不得串用。"),
        ("主商品", "mainFlag商品默认选中且不能取消。", "主商品必须固定选中，并以明确样式区分；未返回主商品时禁止购买并记录数据异常。"),
        ("其他商品", "非主商品可勾选或取消。", "至少再选择1个商品，最终选中商品总数不得少于2个。"),
        ("商品展示", "展示图片、商品名、组合销售价minSalePrice、原价originalPrice和已选规格。", "缺失图片或价格时使用空态，不以后台itemList或Mock值补齐。"),
        ("规格选择", "有规格商品通过SKU弹层选择；查询商品基础信息时携带collocationId，选择后更新skuId、规格和组合价。", "每个选中商品必须具备有效skuId；切换SKU后立即重算套餐价和优惠金额。"),
        ("购买套数", "goodCount默认1，现有组件上限maxCount为999。", "套数作用于全部选中商品；上限应由各商品库存、限购和每组数量共同计算，不能只使用固定999。"),
    ], [1800, 3600, 3960])
    add_body(doc, "现有代码把所有选中SKU的quantity都设置为goodCount。ERP查询接口同时存在itemList.wareqty（每组数量）。若某商品每组数量大于1，正确实际数量应为“每组数量 × 购买套数”；在商城接口未提供或未使用每组数量前，不能认定当前实现支持该类组合。")

    add_heading(doc, "5.7 小程序端-金额计算、购买校验与订单", 2)
    add_table(doc, ["环节", "规则"], [
        ("最高省", "现有逻辑对当前套餐全部商品计算max(原价-originalPrice与组合销售价minSalePrice之差, 0)后求和，用于“最高省X元”展示。"),
        ("套餐价", "仅对当前选中商品的minSalePrice求和，再乘购买套数；按2位小数展示。"),
        ("当前优惠", "仅对当前选中商品的原价与组合销售价差额求和，再乘购买套数；负数按0处理。"),
        ("购买前校验", "选中商品至少2个；每个选中商品必须已选SKU；任一SKU库存小于所需实际数量时禁止进入订单并提示具体商品。"),
        ("订单预览参数", "校验通过后生成collocationId和skuList；skuList包含productId、skuId、quantity，跳转fromPage=collocation。"),
        ("专用预览", "订单页根据fromPage=collocation调用组合购专用预览接口；同时传配送方式，以及地址或自提门店信息。"),
        ("数量锁定", "订单确认页存在collocationId时隐藏普通商品数量加减器，按预览返回数量展示，防止破坏组合结构。"),
        ("权威金额", "弹层金额仅作预估。商品价格、优惠、运费、优惠券资格和应付金额以订单预览接口返回为准。"),
        ("订单提交", "提交时携带预览结果、collocationId、submitToken、配送信息和服务端计算的应付金额，使用统一submitOrder生成订单。"),
    ], [2200, 7160])
    add_body(doc, "商城现有组合购路径为直接进入订单预览，未发现加入普通购物车的组合逻辑。因此本期PRD不把购物车凑单、组合卡片或购物车改数量写成已支持能力。")

    add_heading(doc, "5.8 状态、异常与系统边界", 2)
    add_table(doc, ["场景", "产品处理"], [
        ("后台正常分页", "使用pageNo、pageSize和total分页；切页保留查询条件。startTime包含、endTime不包含。"),
        ("后台空值/空明细", "空字段统一显示“—”；itemList为空时显示“暂无组合商品明细”；cnt与实际条数不一致时记录异常。"),
        ("后台接口错误", "400提示查询条件有误；500或请求失败提示加载失败并保留页面状态。"),
        ("商城组合失效", "入口打开后无可用套餐、SKU失效、库存不足或预览接口判定失效时，不允许按弹层预估价下单。"),
        ("配送变化", "切换骑手配送、快递或自提后重新预览，保留collocationId；若商品不支持当前方式，按预览结果提示。"),
        ("普通购物车", "本期组合购买不加入普通购物车；普通购物车逻辑不得把组合内商品拆成独立优惠行后继续沿用组合价。"),
        ("优惠券/会员价", "订单页存在优惠券入口，但组合价能否叠加优惠券、满减或会员价未在现有组合组件中定义，以预览接口结果为准并列入待确认。"),
        ("售后", "现有代码未发现组合专用退款、换货或优惠回收规则；不得自行假设可按单品原成交价退款。"),
    ], [2200, 7160])
    add_callout(doc, "数据边界", "ERP后台分页接口不返回交易价格、活动时间、状态、门店、图片或库存；商城交易接口会返回组合标题、商品、SKU和价格。两套数据需通过明确的组合标识映射，前端不得按名称猜测关联。")

    add_heading(doc, "5.9 验收标准", 2)
    add_table(doc, ["编号", "验收项", "通过标准"], [
        ("AC-01", "ERP与后台边界", "管理后台仅查询、重置、分页和查看；无新增、编辑、删除、启停、审批或回写ERP。"),
        ("AC-02", "后台字段", "列表与详情仅展示ERP查询接口真实字段；不展示最后修改时间、组合交易价或商品规格等未返回内容。"),
        ("AC-03", "小程序入口", "有可用组合的商品显示可点击入口并能打开组合弹层；无组合、失败或加载中状态均有明确处理。"),
        ("AC-04", "套餐选择", "默认第一组；主商品必选；非主商品可选；切换套餐、SKU或套数后金额与选择状态正确刷新。"),
        ("AC-05", "购买校验", "少于2个商品、缺少SKU、库存不足或每组数量无法正确计算时均不能进入订单。"),
        ("AC-06", "订单预览", "组合购买携带collocationId和完整skuList进入专用预览；配送方式或地址变化后重新计算且保留组合标识。"),
        ("AC-07", "价格一致", "弹层仅展示预估，订单应付以服务端预览为准；预览失败或价格变化时不得沿用旧价格下单。"),
        ("AC-08", "数量完整性", "每个商品实际数量等于每组数量乘购买套数；订单页组合商品数量不可被单独修改。"),
        ("AC-09", "购物车边界", "组合购直接进入订单预览，不把普通购物车组合展示写成已支持功能。"),
        ("AC-10", "安全与显示", "SECRET_KEY不下发前端；管理后台不展示ERP同步、只读、接口原型或Mock数据徽标。"),
    ], [1100, 2300, 5960], center_columns={0})

    add_heading(doc, "5.10 待确认项", 2)
    pending = (
        "高优先级：确认ERP itemList.wareqty如何映射到商城productRespList，以及商城下单quantity是否应为wareqty × goodCount。当前代码统一传goodCount，可能无法支持每组数量大于1。",
        "确认商品详情组合价入口的触发条件和位置。当前组件与打开方法存在，但编译版模板未发现可见触发入口。",
        "确认ERP组合wareid/warecode与商城collocationId的唯一映射方式，以及同步失败、删除、停用和门店切换时的处理。",
        "确认组合价与会员价、优惠券、满减及其他促销的叠加优先级。现有订单页可选择优惠券，但组合规则由预览接口决定。",
        "确认购买套数上限口径：库存、限购、处方药限制及每组数量如何共同确定，替代当前固定maxCount=999。",
        "确认订单行如何保存组合ID、套餐ID、每组数量、购买套数和优惠分摊，以支持对账与订单详情展示。",
        "确认整组退款、部分商品退款、换货和取消后的组合优惠回收及金额分摊规则。",
        "接口示例中cnt与itemList条数不一致，需接口方确认生产数据是否保证一致。",
    )
    for item in pending:
        add_list_item(doc, item, bullet_num_id)

    doc.core_properties.title = "千金小程序需求-组合价"
    doc.core_properties.subject = "ERP组合商品、管理后台查询与小程序组合购"
    doc.core_properties.author = "千金大药房"
    doc.core_properties.keywords = "组合价, 组合购, ERP, 管理后台, 小程序, PRD"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_document()
