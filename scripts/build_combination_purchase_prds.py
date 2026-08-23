from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
ADMIN_OUT = ROOT / "后台系统组合购PRD_V1.2.docx"
MALL_OUT = ROOT / "商城端组合购PRD_V1.2.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
GRAY = "666666"
MUTED = "7A8491"
HEADER_FILL = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
CAUTION = "FFF7E3"
ERROR_FILL = "FFF0F0"
WHITE = "FFFFFF"
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(run, size=11, bold=False, color="000000", italic=False):
    run.font.name = "Calibri"
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), "Calibri")
    rpr.rFonts.set(qn("w:hAnsi"), "Calibri")
    rpr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def configure_style(style, size, color="000000", bold=False, before=0, after=6, line=1.25):
    style.font.name = "Calibri"
    rpr = style._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), "Calibri")
    rpr.rFonts.set(qn("w:hAnsi"), "Calibri")
    rpr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)
    pf = style.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_geometry(table, widths, indent=TABLE_INDENT_DXA):
    if sum(widths) != CONTENT_WIDTH_DXA:
        raise ValueError(f"table widths must total {CONTENT_WIDTH_DXA}, got {sum(widths)}")
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    for tag, value in (("w:tblW", sum(widths)), ("w:tblInd", indent)):
        node = tbl_pr.first_child_found_in(tag)
        if node is None:
            node = OxmlElement(tag)
            tbl_pr.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    for index, width in enumerate(widths):
        table._tbl.tblGrid.gridCol_lst[index].set(qn("w:w"), str(width))
    for row in table.rows:
        prevent_row_split(row)
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table, color="D9DEE7", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def write_cell(cell, text, bold=False, color="000000", size=9.4, align=WD_ALIGN_PARAGRAPH.LEFT):
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.18
    run = paragraph.add_run(str(text))
    set_run_font(run, size=size, bold=bold, color=color)


def add_table(doc, headers, rows, widths, center_columns=None, header_fill=HEADER_FILL):
    center_columns = set(center_columns or [])
    table = doc.add_table(rows=1, cols=len(headers))
    for index, label in enumerate(headers):
        set_cell_shading(table.rows[0].cells[index], header_fill)
        write_cell(
            table.rows[0].cells[index],
            label,
            bold=True,
            color=INK,
            size=9.3,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
    set_repeat_table_header(table.rows[0])
    for row_values in rows:
        if len(row_values) != len(headers):
            raise ValueError("table row length does not match header length")
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            alignment = WD_ALIGN_PARAGRAPH.CENTER if index in center_columns else WD_ALIGN_PARAGRAPH.LEFT
            write_cell(cells[index], value, align=alignment)
    set_table_geometry(table, widths)
    set_table_borders(table)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(2)
    return table


def add_body(doc, text, bold_lead=None):
    paragraph = doc.add_paragraph(style="Normal")
    if bold_lead and text.startswith(bold_lead):
        set_run_font(paragraph.add_run(bold_lead), bold=True, color=INK)
        set_run_font(paragraph.add_run(text[len(bold_lead):]))
    else:
        set_run_font(paragraph.add_run(text))
    return paragraph


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    if level == 1:
        set_run_font(run, size=16, bold=True, color=BLUE)
    elif level == 2:
        set_run_font(run, size=13, bold=True, color=BLUE)
    else:
        set_run_font(run, size=12, bold=True, color=DARK_BLUE)
    return paragraph


def next_num_id(numbering):
    ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    return max(ids, default=0) + 1


def next_abstract_num_id(numbering):
    ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    return max(ids, default=0) + 1


def create_list_numbering(doc, ordered=False):
    numbering = doc.part.numbering_part.element
    abstract_id = next_abstract_num_id(numbering)
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if ordered else "bullet")
    level.append(num_fmt)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1." if ordered else "•")
    level.append(level_text)
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    level.append(justification)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    p_pr.append(indent)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)
    num_id = next_num_id(numbering)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_num_id = OxmlElement("w:abstractNumId")
    abstract_num_id.set(qn("w:val"), str(abstract_id))
    num.append(abstract_num_id)
    numbering.append(num)
    return num_id


def add_list_item(doc, text, num_id):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.extend((ilvl, num_id_node))
    p_pr.append(num_pr)
    set_run_font(paragraph.add_run(text))
    return paragraph


def add_callout(doc, title, text, fill=CALLOUT, title_color=DARK_BLUE):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.08)
    paragraph.paragraph_format.right_indent = Inches(0.08)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.20
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), fill)
    p_pr.append(shading)
    borders = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "4")
        border.set(qn("w:color"), "D5DCE6")
        borders.append(border)
    p_pr.append(borders)
    set_run_font(paragraph.add_run(f"{title}："), size=10.2, bold=True, color=title_color)
    set_run_font(paragraph.add_run(text), size=10.2, color="333333")


def add_meta_row(doc, label, value):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.15
    set_run_font(paragraph.add_run(f"{label}："), size=10.3, bold=True, color=INK)
    set_run_font(paragraph.add_run(value), size=10.3, color="333333")


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, text, end))


def new_document(short_title, title, subtitle, module):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    configure_style(doc.styles["Normal"], 11, after=6, line=1.25)
    configure_style(doc.styles["Heading 1"], 16, BLUE, True, before=18, after=10, line=1.0)
    configure_style(doc.styles["Heading 2"], 13, BLUE, True, before=14, after=7, line=1.0)
    configure_style(doc.styles["Heading 3"], 12, DARK_BLUE, True, before=10, after=5, line=1.0)
    bullet_id = create_list_numbering(doc, ordered=False)
    ordered_id = create_list_numbering(doc, ordered=True)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    set_run_font(header.add_run("千金健康商城 | 产品需求文档"), size=9, color=GRAY)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_after = Pt(0)
    set_run_font(footer.add_run(f"{short_title}  ·  第 "), size=9, color=GRAY)
    add_page_field(footer)
    set_run_font(footer.add_run(" 页"), size=9, color=GRAY)

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(10)
    kicker.paragraph_format.space_after = Pt(3)
    set_run_font(kicker.add_run("PRODUCT REQUIREMENTS DOCUMENT"), size=9.5, bold=True, color=BLUE)

    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    title_p.paragraph_format.keep_with_next = True
    set_run_font(title_p.add_run(title), size=24, bold=True, color=INK)

    subtitle_p = doc.add_paragraph()
    subtitle_p.paragraph_format.space_before = Pt(0)
    subtitle_p.paragraph_format.space_after = Pt(14)
    set_run_font(subtitle_p.add_run(subtitle), size=12.5, color=GRAY)

    add_meta_row(doc, "所属模块", module)
    add_meta_row(doc, "文档版本", "V1.2")
    add_meta_row(doc, "文档日期", "2026-08-23")
    add_meta_row(doc, "目标读者", "产品、运营、交互、前端、后端、订单、售后与测试团队")
    add_meta_row(doc, "业务边界", "商城组合购独立于 ERP；不读取、不同步、不映射、不回写 ERP 组合价数据")
    add_meta_row(doc, "文档状态", "产品口径确认版")
    return doc, bullet_id, ordered_id


def add_shared_business_rules(doc):
    add_table(
        doc,
        ["规则项", "统一口径", "系统约束"],
        [
            ("数据归属", "商城组合购由商城后台独立创建和维护。", "与 ERP 组合价无任何数据同步、ID 映射或状态联动。"),
            ("组合模型", "固定活动总价；不区分主商品。", "交易时失效或库存不足 SKU 从订单移除且组合价格不变；全局具备活动资格的 SKU 少于 2 个时活动自动结束且不恢复，仅单门店库存不足时只隐藏该门店入口。"),
            ("门店范围", "适用于全部门店。", "后台不提供门店选择；结算以最终履约门店为准，地址、定位或门店变化必须重新计价。"),
            ("商品数量", "每个组合至少 2 个、最多 9 个不同商品/SKU。", "商品不可重复；每组数量必须为大于 0 的整数。"),
            ("原价来源", "取最终履约门店实时售价。", "浏览阶段可用用户选择/LBS 门店估算；结算门店变化必须重新计价，不采用会员价或促销价。"),
            ("价格合法性", "创建/编辑时 0 < 组合价格 < 商品原价合计。", "活动期间门店售价变化不修改组合价格；原价合计降至低于组合价时只展示固定组合价，不展示优惠、节省或原价划线。"),
            ("活动状态", "未开始、进行中、已结束。", "未开始和进行中可手动终止；全局具备活动资格的 SKU 少于 2 个时系统自动结束；两种结束均不可恢复且保留计划结束时间。"),
            ("营销互斥", "组合购不参与任何其他价格权益。", "商品只要存在未开始或进行中的组合购/其他促销即冲突，不因时间不重叠而放行；已结束释放占用。"),
            ("购买限制", "不设置组合购专属限购。", "不限制单笔、单用户每日或活动周期套数；库存及商品合规限制仍生效。"),
            ("退款", "含组合购的订单只允许整单退款。", "部分履约或跨仓时以一个售后主单统筹拦截/退回子任务，全部闭环后统一退款。"),
            ("金额分摊", "按实际成交商品的门店原价金额占比进行分摊。", "按分计算并处理尾差，全部成交商品行分摊金额之和必须等于组合成交总额。"),
        ],
        [1500, 3000, 4860],
    )


def build_admin_prd():
    doc, bullets, numbers = new_document(
        "后台组合购 PRD",
        "千金健康商城后台系统｜组合购 PRD",
        "商城独立活动管理 · 固定整组总价 · 全部门店生效",
        "营销 > 平台促销 > 组合购",
    )
    add_callout(
        doc,
        "已确认产品口径",
        "后台组合购与 ERP 独立运作；配置时每组包含 2-9 个商品并设置固定组合价格；商品存在任一未开始/进行中活动即冲突；结束时间必须晚于当前时间；全局具备活动资格的 SKU 少于 2 个时系统自动结束；活动名称和活动单号永久唯一；含组合购订单仅支持整单退款。",
    )

    add_heading(doc, "1. 文档说明", 1)
    add_heading(doc, "1.1 目的", 2)
    add_body(doc, "本文档定义商城后台组合购活动的查询、新增、编辑、终止、商品选择、保存、详情、数据校验与服务端契约，为产品设计、研发实现和测试验收提供统一依据。")
    add_heading(doc, "1.2 本期范围", 2)
    for text in (
        "组合购活动列表、组合条件查询、重置、分页与详情查看。",
        "新增组合购活动，包括活动时间、固定组合价格、2-9 个组合商品、每组数量与描述。",
        "活动创建或编辑成功后写入商城营销服务，并向商城端提供同源查询与交易校验能力。",
        "未开始活动支持编辑和终止，进行中活动支持终止；手动终止记录操作人和时间，全局具备活动资格的 SKU 少于 2 个时系统自动结束。",
    ):
        add_list_item(doc, text, bullets)
    add_heading(doc, "1.3 不在本期范围", 2)
    for text in (
        "ERP 组合价数据接入、同步、映射、回写或冲突处理。",
        "组合购活动复制、删除、恢复已终止活动、审批和批量导入。",
        "指定门店、指定会员、指定渠道或人群定向配置。",
        "主商品、可选搭配商品、单商品组合价或任选商品模式。",
    ):
        add_list_item(doc, text, bullets)

    add_heading(doc, "2. 业务背景与目标", 1)
    add_body(doc, "商城需要独立运营组合购活动，运营人员可将多个商城商品按固定数量组成一个套餐，并设置低于商品原价合计的整组成交价。该能力不再依赖 ERP 的组合价维护或开放接口，商城后台与商城交易服务使用同一套活动主数据。")
    add_heading(doc, "2.1 产品目标", 2)
    for text in (
        "让运营人员在一个页面完成组合购活动创建，并能按活动和商品维度快速查询。",
        "保证后台配置、商城展示、订单预览、成交记录和售后退款使用同一活动版本与价格口径。",
        "通过固定活动总价和失效 SKU 自动剔除/不足 2 个自动结束规则，保证门店商品状态变化后仍有确定口径。",
        "通过服务端强校验避免零价、创建时倒挂价、少于两个商品、重复商品、重名、重号和促销冲突活动进入商城。",
    ):
        add_list_item(doc, text, bullets)

    add_heading(doc, "3. 统一业务规则", 1)
    add_shared_business_rules(doc)
    add_callout(
        doc,
        "价格口径",
        "商品原价取最终履约门店的实时售价，不读取会员价、限时折扣价、券后价或其他促销价。保存活动时记录参考价快照用于校验与审计；活动期间售价变化不调整固定组合价格，即使实时原价合计低于组合价，活动仍继续生效，但商城只展示固定组合价。",
        fill=CAUTION,
        title_color="7A5A00",
    )

    add_heading(doc, "4. 角色与业务流程", 1)
    add_table(
        doc,
        ["步骤", "处理方", "业务动作", "结果"],
        [
            ("1", "运营人员", "进入营销 > 平台促销 > 组合购，点击新增组合购。", "打开空白活动表单。"),
            ("2", "运营人员", "填写名称、开始/结束时间、组合价格与描述。", "形成活动基本信息。"),
            ("3", "运营人员", "从商城商品中心选择 2-9 个商品并设置每组数量。", "形成固定整组商品清单。"),
            ("4", "前端", "计算并展示商品原价合计，执行必填、数量与格式校验。", "阻止明显无效提交。"),
            ("5", "商城营销服务", "校验结束时间晚于当前时间、商品、价格、永久唯一性和未开始/进行中促销占用。", "通过后进入持久化事务。"),
            ("6", "商城营销服务", "生成唯一活动 ID、活动单号和版本号，保存活动及商品快照。", "创建成功并返回活动信息。"),
            ("7", "后台前端", "提示保存成功并返回列表第一页。", "新活动按创建时间倒序显示。"),
            ("8", "商城端", "仅读取进行中且无手动终止/自动结束标记的活动。", "未开始和任一已结束活动均无标识、不生效。"),
            ("9", "运营人员", "对未开始活动编辑任意字段，或对未开始/进行中活动执行终止。", "编辑后版本递增；终止后立即失效。"),
        ],
        [720, 1500, 4560, 2580],
        center_columns={0},
    )

    add_heading(doc, "5. 功能需求", 1)
    add_heading(doc, "5.1 功能入口与权限", 2)
    add_table(
        doc,
        ["项目", "要求"],
        [
            ("入口", "营销 > 平台促销 > 组合购。功能入口、列表标题和详情面包屑统一使用“组合购”。"),
            ("可用操作", "查询、重置、分页、新增、查看；未开始支持编辑/终止，进行中支持终止，已结束仅查看。"),
            ("权限", "查看、新增、编辑和终止权限分别由服务端校验；不得仅依赖前端按钮显隐。"),
            ("数据范围", "活动固定适用全部门店，不提供门店选择或门店数据权限筛选。"),
            ("禁用操作", "本期不提供复制、删除、恢复已终止活动、审批或 ERP 操作。"),
        ],
        [2100, 7260],
    )

    add_heading(doc, "5.2 组合购列表", 2)
    add_heading(doc, "5.2.1 查询条件", 3)
    add_table(
        doc,
        ["查询项", "匹配方式", "长度/格式", "规则"],
        [
            ("活动单号", "精确", "最多 50 字符", "去除首尾空格后精确匹配；保留前导 0。"),
            ("活动名称", "模糊", "最多 100 字符", "名称包含匹配；去除首尾空格。"),
            ("活动起止时间", "区间重叠", "开始日期-结束日期", "活动结束时间 >= 查询开始时间，且活动开始时间 <= 查询结束时间；查询结束日期包含当日 23:59:59。"),
            ("商品名称", "模糊", "最多 100 字符", "任一组合商品名称包含输入值即命中活动。"),
            ("商品编码", "精确", "最多 50 字符", "任一组合商品编码精确匹配即命中活动。"),
        ],
        [1650, 1350, 1800, 4560],
    )
    add_body(doc, "查询同时应用全部已填写条件，各条件之间为 AND；空白条件表示不限制。点击查询后回到第 1 页。点击重置清空全部条件、日期、页码和跳转值，并恢复默认列表。")
    add_heading(doc, "5.2.2 列表字段与分页", 3)
    add_table(
        doc,
        ["列", "展示规则"],
        [
            ("活动单号", "展示服务端生成的活动单号，按字符串处理。"),
            ("活动名称", "完整展示或省略显示，悬停可查看完整名称。"),
            ("活动时间", "格式为 YYYY-MM-DD HH:mm:ss ~ YYYY-MM-DD HH:mm:ss。"),
            ("活动状态", "位于活动时间之后；时间自然结束、手动终止或系统自动结束均显示为已结束。"),
            ("操作", "未开始：查看、编辑、终止；进行中：查看、终止；已结束：仅查看。"),
        ],
        [2100, 7260],
    )
    add_body(doc, "列表不展示营销类型、商品明细、商品数、规格/单位、厂家、组合价格或门店列。默认按创建时间倒序；默认 10 条/页，可切换 10、20、50、100 条/页。")
    add_callout(doc, "状态公式", "存在手动终止时间 terminatedAt 或系统自动结束时间 autoEndedAt 时状态直接为已结束；否则当前中国标准时间 < 开始时间为未开始，开始时间 <= 当前时间 <= 结束时间为进行中，当前时间 > 结束时间为已结束。结束标记覆盖时间计算，但不得覆盖原计划结束时间。")

    add_heading(doc, "5.3 新增组合购", 2)
    add_heading(doc, "5.3.1 表单字段", 3)
    add_table(
        doc,
        ["字段", "必填", "控件/限制", "业务规则"],
        [
            ("活动名称", "是", "文本框，最多 100 字符", "默认空，提示“请输入活动名称”；去除首尾空格后全局唯一，已结束活动也参与校验。"),
            ("开始时间", "是", "日期时间选择器，精确到秒", "允许早于当前时间，但不得晚于已选结束时间；空态“请选择开始时间”。"),
            ("结束时间", "是", "日期时间选择器，精确到秒", "必须同时严格晚于开始时间和保存时的当前中国标准时间；空态“请选择结束时间”。"),
            ("组合价格", "是", "金额输入，2 位小数", "必须大于 0，且严格低于商品原价合计。"),
            ("组合商品", "是", "商品选择弹窗 + 明细表", "2-9 个不同商品/SKU；不区分主商品；全部商品为必购。"),
            ("组合描述", "否", "多行文本，最多 500 字符", "默认空；实时展示已输入字符数。"),
            ("适用门店", "系统固定", "不提供可编辑控件", "服务端统一保存为 ALL，含义为全部门店。"),
        ],
        [1500, 900, 2850, 4110],
        center_columns={1},
    )
    add_heading(doc, "5.3.2 时间选择", 3)
    for text in (
        "开始时间选择器禁用已选结束日期之后的日期；同日时，开始时间不得晚于结束时间。",
        "结束时间选择器禁用已选开始日期和当前日期两者中较晚者之前的日期；同日时，下限取开始时间与当前时间的较晚值。",
        "保存时服务端仍须校验结束时间严格晚于开始时间和当前中国标准时间，不能只依赖前端控件。",
        "开始时间允许早于保存时刻；若开始时间 <= 当前时间 < 结束时间，保存后立即进入进行中；不允许创建即结束。",
        "活动时间按中国标准时间保存和计算，接口使用带时区的标准时间或明确约定 +08:00。",
    ):
        add_list_item(doc, text, bullets)

    add_heading(doc, "5.3.3 前端校验与提示", 3)
    add_table(
        doc,
        ["场景", "页面处理", "提示文案"],
        [
            ("活动名称为空", "字段红框，错误紧邻控件显示。", "请输入活动名称"),
            ("开始时间为空", "字段红框，错误紧邻控件显示。", "请选择开始时间"),
            ("结束时间为空", "字段红框，错误紧邻控件显示。", "请选择结束时间"),
            ("组合价格为空", "字段红框，错误紧邻控件显示。", "请输入组合价格"),
            ("组合价格 <= 0", "阻止保存并定位金额字段。", "组合价格必须大于0"),
            ("组合价格 >= 原价合计", "阻止保存并展示当前原价合计。", "组合价格必须低于商品原价合计"),
            ("未选择商品", "商品控件不加红框，页面上方 Toast。", "请添加组合商品"),
            ("只选择 1 个商品", "页面上方 Toast。", "组合商品至少选择2个"),
            ("结束时间 <= 开始时间", "页面上方 Toast，不在保存按钮旁显示。", "结束时间必须晚于开始时间"),
            ("结束时间 <= 当前时间", "页面上方 Toast，新增和编辑均阻止保存。", "结束时间必须晚于当前时间"),
            ("第 10 个商品", "本次选择不生效，原选择保持不变。", "组合商品最多只能选择9个~"),
            ("活动名称重复", "阻止保存；已结束活动名称同样不可复用。", "活动名称已存在"),
            ("商品促销冲突", "只要商品存在任一未开始/进行中活动就阻止选择或保存，不判断时间是否重叠。", "商品已参与其他组合购/其他促销"),
        ],
        [2400, 4110, 2850],
    )
    add_body(doc, "首次点击保存后，活动名称、开始时间、结束时间和组合价格分别显示自己的红色错误；每个字段完成后只清除该字段错误。前四个必填行预留错误文案高度，避免页面跳动。")

    add_heading(doc, "5.4 商品选择弹窗", 2)
    add_table(
        doc,
        ["区域", "要求"],
        [
            ("弹窗", "使用大尺寸商品选择弹窗，标题为“商品选择”，支持关闭、取消和确认。"),
            ("分类筛选", "14 px 触发器默认显示“请选择商品分类”；打开后左侧一级分类、右侧当前悬停一级分类的二级分类。"),
            ("名称/编码", "14 px 组合查询控件；商品名称按包含匹配，完整商品编码按精确匹配。"),
            ("表格", "列为商品编码、商品名称、最新原价；商品编码列含行复选框并保持单行显示。"),
            ("本页选择", "表头不设置全选复选框；在表格工具区提供“选择本页/取消本页”，仅影响当前页且仍受 9 个上限约束。"),
            ("跨页状态", "翻页、搜索或切换分类后保留已选商品；弹窗再次打开时按表单当前商品初始化。"),
            ("反选同步", "已满 9 个时仍可取消已有商品；确认后从表单移除被取消商品，并保留所有继续选中商品的数量。"),
            ("分页", "支持 10、20、50 条/页、上一页、下一页、页码和跳转。"),
            ("确认", "只有选择发生变化时确认按钮可用；确认后以弹窗勾选结果完整同步表单。"),
            ("促销占用", "已被未开始/进行中组合购或其他促销占用的商品置灰禁选；已结束活动不再占用。"),
        ],
        [2100, 7260],
    )

    add_heading(doc, "5.5 已选组合商品", 2)
    add_table(
        doc,
        ["列", "展示/交互规则"],
        [
            ("商品图片", "展示商品主图；缺失时使用商城默认图。"),
            ("商品编码", "展示商城商品/SKU编码并保持单行。"),
            ("商品名称", "展示完整名称或省略，悬停显示完整内容。"),
            ("最新原价", "展示商城普通原价，保留 2 位小数；价格变化时重新计算原价合计。"),
            ("商品规格", "展示活动绑定的具体 SKU 规格。"),
            ("数量", "居中边框输入框，默认 1，只允许大于 0 的整数；标题与输入均居中。"),
            ("操作", "仅提供删除，标题与删除动作均居中。"),
        ],
        [2100, 7260],
    )
    add_body(doc, "明细表下方固定显示“组合商品最多可添加9个”。不展示组合单品价、近效期字段或主商品标识。商品数量修改后立即更新商品原价合计。")

    add_heading(doc, "5.6 保存与幂等", 2)
    for text in (
        "点击保存后立即禁用保存按钮并显示处理中，避免重复提交。",
        "客户端生成 requestId；服务端按 requestId 幂等，同一请求重复到达只创建一条活动。",
        "服务端在一个事务中保存活动主表、商品明细、原价快照和版本号；任一明细失败则整体回滚。",
        "保存成功提示“保存成功”，返回列表第 1 页并将新活动置顶；保存失败保留已填内容和商品数量。",
        "活动单号由服务端生成，最长 50 字符；生成后永久不可编辑、删除后复用或由前端拼接。",
        "活动名称去除首尾空格后与全部历史活动比较；活动名称和活动单号均建立唯一约束，且该约束永久有效，服务端在事务内执行最终校验。",
        "保存新增/编辑前校验全部商品促销占用；只要其他活动为未开始或进行中即冲突，不判断活动时间是否重叠；已结束活动不冲突。",
    ):
        add_list_item(doc, text, bullets)

    add_heading(doc, "5.7 组合购详情", 2)
    add_table(
        doc,
        ["区域", "字段"],
        [
            ("活动基本信息", "活动单号、活动名称、活动状态、开始时间、计划结束时间、适用门店（全部门店）、组合价格、商品原价合计、组合描述、创建人、创建时间。"),
            ("商品明细", "图片、商品编码、商品名称、规格、创建时原价快照、每组数量、原价小计。"),
            ("结束信息", "手动终止展示终止时间和操作人；系统自动结束展示结束时间和原因。计划结束时间保持原值，不被覆盖。"),
            ("操作", "仅提供返回列表；返回后恢复此前查询条件、页码和每页条数。"),
        ],
        [2100, 7260],
    )
    add_body(doc, "详情复用新增表单的字段结构和商品明细表，但所有控件只读，不提供商品选择、数量修改、删除或保存。详情不展示 ERP 字段、营销类型、组合厂家、组合单位或主商品标记。")

    add_heading(doc, "5.8 编辑与终止", 2)
    add_table(
        doc,
        ["能力", "可操作状态", "产品与服务端规则"],
        [
            ("编辑", "仅未开始", "打开与新增相同的完整表单，所有活动字段、商品和数量均回填且可编辑；保存时重复执行唯一性、时间、价格、数量和促销冲突校验，成功后版本号递增。"),
            ("手动终止", "未开始、进行中", "二次确认后记录 terminatedAt 和 terminatedBy，并在数据库操作日志写入操作人、终止时间；状态立即变为已结束。"),
            ("系统自动结束", "进行中", "全局下架、停售或删除后具备活动资格的 SKU 少于 2 个时，原子写入 autoEndedAt 和原因 SELLABLE_SKU_LT_2，立即释放促销占用且不自动恢复；单门店库存不足不触发全局结束。"),
            ("已结束", "仅查看", "自然结束、手动终止或系统自动结束后均不可编辑、再次终止或恢复；历史订单继续使用成交快照。"),
        ],
        [1500, 1800, 6060],
    )
    add_callout(doc, "结束语义", "手动终止和系统自动结束均为不可逆生命周期操作，不删除活动、不覆盖计划结束时间，也不影响已生成订单。终止前已创建但未支付订单可继续支付至订单自身超时。权限、状态和并发版本必须由服务端再次校验；重复请求按幂等成功返回。", fill=CAUTION, title_color="7A5A00")

    add_heading(doc, "5.9 空态与异常", 2)
    add_table(
        doc,
        ["场景", "处理"],
        [
            ("查询无结果", "列表区域显示“暂无符合条件的组合购活动”，保留查询条件。"),
            ("商品查询无结果", "弹窗显示“暂无符合条件的商品”，已选跨页状态不丢失。"),
            ("商品已下架/删除", "保存前阻止创建并定位对应商品，提示“商品已失效，请重新选择”。"),
            ("全局活动资格 SKU 少于 2 个", "系统将活动永久自动结束，记录结束时间和原因，释放促销占用；SKU 恢复后活动不自动恢复。单门店库存不足只隐藏该门店入口。"),
            ("商品原价变化", "新增/未开始编辑保存时刷新参考原价并校验；活动进行中门店售价下降至原价合计低于组合价时，活动和组合价格保持不变。"),
            ("保存接口失败", "提示“保存失败，请稍后重试”，表单内容保持不变。"),
            ("结束并发冲突", "重新拉取活动状态；已被手动/自动结束时按幂等成功处理，已自然结束时刷新为已结束。"),
            ("列表/详情接口失败", "显示加载失败和重试入口，不清空查询条件。"),
        ],
        [2700, 6660],
    )

    add_heading(doc, "6. 数据与接口契约", 1)
    add_heading(doc, "6.1 活动主数据", 2)
    add_table(
        doc,
        ["字段", "类型", "必填", "说明"],
        [
            ("activityId", "string", "是", "商城组合购内部唯一 ID，与 ERP 无关。"),
            ("activityNo", "string(50)", "是", "服务端生成后不可修改且永久唯一；任何结束状态或逻辑删除后均不可复用。"),
            ("activityName", "string(100)", "是", "去除首尾空格后永久唯一；所有历史状态均参与比较。"),
            ("startTime/endTime", "datetime(+08:00)", "是", "有效区间；新增/编辑时 endTime 必须严格晚于服务端当前时间。"),
            ("bundlePrice", "decimal(12,2)", "是", "每套固定组合成交价，必须 >0 且 < 原价合计。"),
            ("storeScope", "enum", "是", "固定为 ALL。"),
            ("description", "string(500)", "否", "活动描述。"),
            ("version", "integer", "是", "交易校验和订单快照使用的活动版本。"),
            ("terminatedAt/terminatedBy", "datetime/string", "否", "手动终止时间和操作人；存在 terminatedAt 时状态强制为已结束。"),
            ("autoEndedAt/autoEndReason", "datetime/enum", "否", "全局活动资格 SKU 少于 2 个等系统自动结束信息；存在 autoEndedAt 时状态强制为已结束。"),
            ("createdBy/createdAt", "string/datetime", "是", "创建审计信息。"),
            ("updatedBy/updatedAt", "string/datetime", "否", "最近一次编辑审计信息。"),
        ],
        [1800, 1800, 900, 4860],
        center_columns={2},
    )
    add_heading(doc, "6.2 商品明细数据", 2)
    add_table(
        doc,
        ["字段", "类型", "必填", "说明"],
        [
            ("productId/skuId", "string", "是", "商城商品与具体可交易 SKU 标识。"),
            ("productCode", "string(50)", "是", "商品/SKU编码。"),
            ("productName", "string(100)", "是", "商品名称快照。"),
            ("specification", "string", "是", "SKU规格快照。"),
            ("imageUrl", "string", "否", "商品图片快照。"),
            ("originalUnitPriceSnapshot", "decimal(12,2)", "是", "创建时普通原价快照，仅用于校验与审计。"),
            ("quantityPerBundle", "integer", "是", "每套所需数量，必须 >=1。"),
            ("sortOrder", "integer", "是", "商品展示顺序。"),
        ],
        [2250, 1800, 900, 4410],
        center_columns={2},
    )
    add_heading(doc, "6.3 逻辑接口", 2)
    add_table(
        doc,
        ["接口", "用途", "关键要求"],
        [
            ("组合购分页查询", "后台列表与筛选", "服务端完成组合商品 EXISTS 查询、时间重叠查询和分页。"),
            ("组合购详情查询", "后台详情", "返回活动主数据、全部商品明细和创建快照。"),
            ("新增组合购", "创建活动", "支持 requestId 幂等；服务端重复执行全部业务校验。"),
            ("编辑组合购", "更新未开始活动", "携带 activityId + version；只允许未开始状态，更新成功后版本递增。"),
            ("终止组合购", "手动结束活动", "只允许未开始或进行中；记录终止审计并立即对商城失效，接口须幂等。"),
            ("自动结束活动", "商品全局资格处理", "全局活动资格 SKU 少于 2 个时原子写结束标记与原因、释放占用、清理缓存且不自动恢复；单店库存不足不调用该接口。"),
            ("促销冲突校验", "商品选择与保存", "校验所有未开始/进行中的组合购及其他促销占用，不判断时间重叠；已结束不冲突。"),
            ("商城商品搜索", "商品选择弹窗", "返回分类、编码、名称、SKU、规格、图片和当前普通原价。"),
            ("商城有效组合查询", "商城端展示与交易", "只返回进行中且 terminatedAt/autoEndedAt 均为空的数据；禁止调用 ERP 组合价接口。"),
        ],
        [2400, 2700, 4260],
    )
    add_callout(doc, "接口命名说明", "本 PRD 只定义逻辑能力和业务字段，实际 URL、鉴权方式、错误码与数据库表名由技术方案确定；任何接口都不得依赖 ERP 组合价 ID 或同步状态。")

    add_heading(doc, "7. 非功能要求", 1)
    add_table(
        doc,
        ["类别", "要求"],
        [
            ("一致性", "活动主表与商品明细事务保存；商城交易按 activityId + version 校验，订单保存完整快照。"),
            ("安全", "查看、新增、编辑和终止权限分别在服务端校验；手动终止的数据库操作日志必须记录操作人、终止时间和变更快照。"),
            ("幂等", "新增、编辑和终止接口均支持 requestId 幂等；同一请求不得重复创建、重复更新版本或重复写终止记录。"),
            ("时间", "服务端统一使用中国标准时间判断状态，客户端时间仅用于展示。"),
            ("性能", "列表与详情正常负载下 P95 不高于 2 秒；创建接口 P95 不高于 3 秒。"),
            ("可观测性", "记录保存失败、唯一性/促销冲突、商品失效、价格校验、手动终止、自动结束和接口异常日志，支持 activityId/requestId 检索。"),
        ],
        [1800, 7560],
    )

    add_heading(doc, "8. 验收标准", 1)
    add_table(
        doc,
        ["编号", "验收项", "通过标准"],
        [
            ("AC-01", "ERP 独立", "后台活动不读取、不展示、不映射和不回写任何 ERP 组合价数据。"),
            ("AC-02", "列表查询", "五个查询条件按规定方式同时生效；重置清空全部条件。"),
            ("AC-03", "时间重叠", "活动区间与查询区间有任意交集即返回，结束日期包含当日。"),
            ("AC-04", "实时状态", "自然结束、手动终止和系统自动结束均正确显示已结束，且不覆盖计划结束时间。"),
            ("AC-05", "商品数量", "少于 2 个或超过 9 个商品均不能保存；重复商品不能加入。"),
            ("AC-06", "固定整组", "表单和数据模型均无主商品、可选商品或单品组合价字段。"),
            ("AC-07", "价格校验", "组合价格为空、<=0 或 >= 保存时参考原价合计时均不能新增或编辑；活动期间售价下降不自动终止活动。"),
            ("AC-08", "商品弹窗", "分类、名称/编码、跨页选择、重新打开初始化、反选同步和 9 个限制均正确。"),
            ("AC-09", "数量保留", "重新打开商品弹窗并确认后，继续选中商品的每组数量不丢失。"),
            ("AC-10", "表单错误", "四个核心必填字段显示各自红框和错误文案，完成后分别清除。"),
            ("AC-11", "时间错误", "结束时间不晚于开始时间时显示上方 Toast，活动不创建。"),
            ("AC-12", "全部门店", "创建请求固定保存 storeScope=ALL，页面没有可配置门店范围。"),
            ("AC-13", "幂等保存", "重复点击或重试同一 requestId 只产生一条活动。"),
            ("AC-14", "详情一致", "详情完整展示新建时的时间、价格、描述和商品数量，不出现旧 ERP 字段。"),
            ("AC-15", "同源可用", "创建成功后商城可从商城营销服务读取活动，无需等待 ERP 同步。"),
            ("AC-16", "编辑权限", "只有未开始活动显示编辑；表单完整回填且可修改全部字段，保存后版本递增。"),
            ("AC-17", "终止权限", "未开始和进行中活动均显示终止；确认后立即变为已结束，商城不再展示并拒绝新订单。"),
            ("AC-18", "已结束操作", "自然结束、手动终止和系统自动结束活动均只显示查看，不可编辑、再次终止或恢复。"),
            ("AC-19", "全局唯一", "活动名称和活动单号在全部历史状态中均不可重复，并发保存由数据库唯一约束兜底。"),
            ("AC-20", "促销冲突", "商品存在任一未开始/进行中的组合购或其他促销即冲突，不因时间不重叠放行；活动结束后释放。"),
            ("AC-21", "活动时间", "开始时间可早于当前时间，但结束时间必须晚于开始时间和当前时间；保存后可立即进入进行中但不可创建即结束。"),
            ("AC-22", "自动结束", "可售 SKU 从 2 个降为 1 个或 0 个时活动永久自动结束、立即释放占用且商品恢复后不自动恢复。"),
            ("AC-23", "终止日志", "手动终止的数据库操作日志可按 activityId 查询到操作人和终止时间，重复请求不重复写日志。"),
        ],
        [1050, 2100, 6210],
        center_columns={0},
    )

    add_heading(doc, "9. 技术实施注意事项", 1)
    for text in (
        "上线前由技术方案确定活动单号生成规则、接口 URL、错误码、鉴权方式和数据库索引。",
        "后台参考原价快照只用于新增/编辑校验与审计；商城成交和分摊必须改用最终履约门店的实时售价。",
        "手动终止与自动结束需要低延迟失效链路：清理或短 TTL 缓存，并在预览、提交校验 terminatedAt/autoEndedAt。",
        "活动名称、活动单号永久唯一索引和未结束商品促销占用由数据库约束或事务锁兜底，冲突判断不得加入时间重叠豁免。",
        "历史旧 ERP 组合价页面或文档不得复用商城组合购字段，避免形成隐性同步关系。",
    ):
        add_list_item(doc, text, bullets)

    doc.core_properties.title = "千金健康商城后台系统组合购PRD"
    doc.core_properties.subject = "商城独立组合购活动管理"
    doc.core_properties.author = "千金健康商城产品团队"
    doc.core_properties.keywords = "组合购, 后台, 固定总价, 全部门店, PRD"
    doc.save(ADMIN_OUT)


def build_mall_prd():
    doc, bullets, numbers = new_document(
        "商城组合购 PRD",
        "千金健康商城端｜组合购 PRD",
        "固定组合价 · 当前门店实时售价 · 营销互斥 · 整单退款",
        "商城商品详情、组合购选择、订单预览、订单与售后",
    )
    add_callout(
        doc,
        "已确认产品口径",
        "商城组合购与 ERP 完全独立；每套使用固定组合价格且适用全部门店；结算以最终履约门店实时售价为原价；不可售 SKU 被移除但组合价格不变，全局活动资格 SKU 少于 2 个时活动永久自动结束，单店库存不足只隐藏该店入口；原价不高于组合价时只展示固定组合价；只要订单含组合购就只允许整单退款。",
    )

    add_heading(doc, "1. 文档说明", 1)
    add_heading(doc, "1.1 目的", 2)
    add_body(doc, "本文档定义消费者在千金健康商城发现、选择、预览、提交和查看组合购订单，以及发起整单退款的完整体验，并明确生命周期、门店实时售价、动态可交易商品集合、活动互斥、金额分摊和异常售后。")
    add_heading(doc, "1.2 本期范围", 2)
    for text in (
        "商品详情页组合购入口、可用组合列表、组合详情、加入购物车和购买套数选择。",
        "按当前门店实时计算可交易商品集合、组合订单预览和提交。",
        "固定组合价格展示、商品行成交金额分摊、订单详情和整单退款。",
        "未开始隐藏、手动终止、全局活动资格 SKU 少于 2 个自动结束、单店库存不足隐藏、履约门店切换、商品失效和原价变化等异常处理。",
    ):
        add_list_item(doc, text, bullets)
    add_heading(doc, "1.3 不在本期范围", 2)
    for text in (
        "ERP 组合价活动展示、同步、关联或交易。",
        "普通散品凑单后自动识别为组合购或跨订单凑单。",
        "主商品 + 可选搭配、消费者主动任选、替换商品或拆分组合价格。",
        "会员价比较、优惠券、折扣、满减、积分抵现及其他营销叠加。",
        "部分商品退款、部分数量退款、部分套数退款和通过退款接口处理单品异常。",
    ):
        add_list_item(doc, text, bullets)

    add_heading(doc, "2. 业务目标与成功口径", 1)
    add_body(doc, "组合购的目标是以固定活动价格提升关联商品连带购买。活动配置包含 2-9 个商品；交易时系统按最终履约门店剔除不可售 SKU，消费者不能自行选择或替换。当前门店少于 2 个可交易 SKU 时只隐藏该店入口；全局下架、停售或删除后活动资格 SKU 少于 2 个时系统永久自动结束活动。")
    add_heading(doc, "2.1 产品目标", 2)
    for text in (
        "消费者能从任一组合成员商品的详情页发现当前门店可购买的组合。",
        "消费者无需凑单，确认系统计算出的当前可交易商品集合与套数后直接进入组合订单预览。",
        "商城展示金额、订单预览金额、支付金额、订单明细和整单退款金额口径一致。",
        "服务端在预览和提交时校验自然时间、手动终止/自动结束标记、最终履约门店、实际商品集合和库存；原价下降不改固定组合价格。",
    ):
        add_list_item(doc, text, bullets)

    add_heading(doc, "3. 统一业务规则", 1)
    add_shared_business_rules(doc)
    add_callout(doc, "交易模型", "一次组合购交易选择一种组合活动，可购买一套或多套；支持加入购物车和立即购买。每点击一次加入购物车增加 1 套完整组合，同一活动版本与履约门店的组合合并并累加套数；不设置活动级单笔、单用户每日或活动周期累计套数上限，但仍受真实库存、禁限售及普通合规规则约束。只要订单含组合购标识，售后均锁定为整单退款。", fill=CAUTION, title_color="7A5A00")

    add_heading(doc, "4. 用户流程", 1)
    add_table(
        doc,
        ["步骤", "页面/系统", "用户或系统动作", "结果"],
        [
            ("1", "商品详情", "商城根据当前商品/SKU和当前门店查询可用组合购。", "有可用组合时展示组合购入口。"),
            ("2", "组合购面板", "系统按当前门店展示活动名称、固定组合价格、实时原价合计及本次可交易商品。", "用户理解实际成交清单；不可售商品不展示。"),
            ("3", "组合购面板", "如存在多个组合，用户切换目标组合；点击一次加入购物车增加 1 套，或点击立即购买进入结算。", "形成 activityId、version、storeId 和 setCount。"),
            ("4", "资格校验", "系统确认活动进行中且无结束标记，并按当前履约门店计算至少 2 个可交易 SKU。", "当前门店少于 2 个时隐藏该店入口；仅全局活动资格少于 2 个时原子结束活动。"),
            ("5", "订单预览", "服务端重新计算实际商品集合、门店实时原价、固定组合价、分摊、配送费与应付金额。", "返回预览快照、商品集合摘要和 previewToken。"),
            ("6", "订单预览", "用户确认收货地址/自提门店和配送方式。", "每次变化均重新预览。"),
            ("7", "订单提交", "用户提交订单，服务端用 previewToken 和幂等键校验并锁定库存。", "生成组合购订单。"),
            ("8", "订单/售后", "用户查看组合购标识、价格与商品分摊；含组合购订单如需退款只能整单申请。", "整单退款或通过补发、客服补偿处理单品异常。"),
        ],
        [720, 1440, 4740, 2460],
        center_columns={0},
    )

    add_heading(doc, "5. 功能需求", 1)
    add_heading(doc, "5.1 商品详情入口", 2)
    add_table(
        doc,
        ["场景", "产品规则"],
        [
            ("入口显示", "当前所选 SKU 命中进行中、无手动终止/自动结束标记且版本有效的活动，并且当前履约门店至少有 2 个活动商品可交易时展示入口。"),
            ("入口隐藏", "活动未开始时不展示入口、角标或任何活动标识；自然结束、手动终止、系统自动结束、可交易商品少于 2 个或数据异常时隐藏。"),
            ("多个组合", "入口展示可用组合数量或最低组合价格；点击后打开组合列表，默认选择推荐顺序第一项。"),
            ("门店来源", "浏览阶段使用用户当前选择或 LBS 定位门店；订单结算以地址/自提方式最终确定的履约门店为准。"),
            ("门店切换", "地址、定位、手动门店或履约门店变化后立即重新查询商品、售价和库存并重新计价；上一门店 previewToken 失效。"),
            ("加载失败", "不影响普通购买；提示“组合购加载失败，请稍后重试”，不得展示缓存旧价作为可成交价。"),
        ],
        [2400, 6960],
    )

    add_heading(doc, "5.2 组合购选择面板", 2)
    add_table(
        doc,
        ["区域", "展示与交互"],
        [
            ("活动信息", "原价合计高于组合价时可展示门店原价合计、固定组合价和节省金额；原价合计等于或低于组合价时只展示固定组合价，不展示优惠、节省或原价划线。"),
            ("组合商品", "只展示当前门店实际可交易的商品图片、名称、固定 SKU 规格、每套数量和实时售价；无勾选框、删除或替换操作。"),
            ("动态剔除", "配置商品下架、停售、删除或库存不足一套时不展示该商品；组合价格保持不变。当前门店剩余至少 2 个 SKU 时继续成交，少于 2 个时仅隐藏该店入口；全局活动资格 SKU 少于 2 个时活动永久自动结束并释放占用。"),
            ("固定 SKU", "活动绑定具体 SKU；消费者不能切换到活动外规格，也不能主动恢复被系统剔除的商品。"),
            ("购买套数", "默认 1；用户调整套数时全部实际商品数量同步变化，实际数量 = 每套数量 × 套数。"),
            ("套数约束", "不设置组合购活动级上限；页面可输入的最大套数仅受当前实际商品库存、禁限售和普通合规规则约束。"),
            ("操作", "提供“加入购物车”和“立即购买”。每点击一次“加入购物车”增加 1 套完整组合，购物车角标按组合套数加 1；同一活动版本与履约门店的组合合并为一行。"),
            ("购物车整组约束", "加入前按增加后的总套数重新校验活动、门店、实际商品集合和库存；失败时数量不变并提示原因。购物车中的组合商品整组绑定，不允许只勾选、删除或修改部分商品。"),
        ],
        [2100, 7260],
    )
    add_body(doc, "配置阶段固定 2-9 个商品及数量，不区分主商品；交易阶段只允许系统按商品状态和库存剔除不可交易 SKU，消费者不能任选。页面不得出现“主商品”“至少再选一件”或单商品组合售价等旧逻辑。")

    add_heading(doc, "5.3 可交易资格", 2)
    add_table(
        doc,
        ["校验项", "通过条件", "失败处理"],
        [
            ("活动", "当前中国标准时间处于活动区间内，terminatedAt/autoEndedAt 均为空且版本有效。", "未开始不展示；任一结束状态关闭入口并拒绝预览/提交。"),
            ("门店", "活动范围为全部门店，结算 storeId 为地址/自提方式确定的最终履约门店。", "门店未确定或变化时重新查询、计价并使旧令牌失效。"),
            ("商品", "从配置清单剔除不可售 SKU 后，当前门店至少剩余 2 个可交易 SKU。", "当前门店少于 2 个时隐藏入口并阻止购买；全局活动资格少于 2 个时原子写入系统自动结束标记。"),
            ("库存", "实际商品均满足一套基础数量，所选套数不超过实际商品库存可支持值。", "不足一套的 SKU 被剔除；多套库存不足时要求减少套数或重新预览。"),
            ("价格", "组合价格 >0；实时读取最终履约门店实际商品售价作为分摊原价。", "原价合计降至组合价以下不影响活动或固定价格，前端只展示固定组合价。"),
            ("分摊基数", "实际商品订单原价合计 Oₙ >0。", "若全部实际商品门店售价均为 0，活动本身不终止，但当前门店阻止预览并告警；不得擅自改用均分。"),
            ("营销", "活动商品不存在任何未开始/进行中的其他组合购或促销，本交易仅使用组合购价格。", "冲突不判断活动时间是否重叠；发现异常时拒绝成交并告警。"),
        ],
        [1800, 4260, 3300],
    )

    add_heading(doc, "5.4 价格计算", 2)
    add_body(doc, "定义：S 为最终履约门店中本次实际可交易商品集合且 |S|>=2；pᵢ 为 S 中商品 i 的履约门店实时普通售价，qᵢ 为每套数量，n 为购买套数，B 为固定每套组合价格。所有金额按人民币计算。")
    add_table(
        doc,
        ["金额", "公式", "说明"],
        [
            ("每套原价合计 O", "Σᵢ∈S(pᵢ × qᵢ)", "只计算实际可交易商品；不使用会员价、促销价或券后价。"),
            ("订单原价合计 Oₙ", "O × n", "本次订单内实际组合商品的门店实时原价合计。"),
            ("组合成交总额 T", "B × n", "不含运费及非商品服务费。"),
            ("价差 D", "Oₙ - T", "可为 0 或负数；D<=0 时只展示固定组合价，不展示 Oₙ、优惠、节省或原价划线。"),
            ("商品行原价金额 Oᵢ", "pᵢ × qᵢ × n", "用于分摊权重和订单快照。"),
        ],
        [2250, 2250, 4860],
    )
    add_callout(doc, "权威金额", "实际商品集合、最终履约门店原价、组合成交总额、库存、配送费和应付金额以服务端订单预览为准。地址、定位或门店变化必须重新计价。活动期间原价调整不修改 B，也不触发 B<O 的交易拦截。")

    add_heading(doc, "5.5 商品行金额分摊", 2)
    add_body(doc, "组合成交总额按本次实际成交商品行的门店原价金额占订单原价合计的比例分摊。原价金额越高，分摊越多。被系统剔除的 SKU 不参与分母和分摊。分摊仅用于订单行、财务、发票和整单退款核对，不代表商品可单独退款。")
    add_table(
        doc,
        ["步骤", "计算规则"],
        [
            ("1. 计算权重", "wᵢ = Oᵢ / Oₙ。"),
            ("2. 计算理论分摊", "Aᵢ* = T × wᵢ。"),
            ("3. 按分截断", "每个商品行先向下截断到 0.01 元，得到 Aᵢ。"),
            ("4. 分配尾差", "剩余分数按 Aᵢ* 的小数余数从大到小每次加 0.01 元；余数相同按商品行原价金额从高到低，再按 SKU ID 升序。"),
            ("5. 校验", "ΣAᵢ 必须严格等于 T；不相等时订单预览失败，不允许提交。"),
        ],
        [2250, 7110],
    )
    add_heading(doc, "5.5.1 分摊示例", 3)
    add_body(doc, "某组合每套包含 A、B、C 三个商品，原价金额分别为 60 元、30 元和 10 元，原价合计 100 元，组合价格 80 元。")
    add_table(
        doc,
        ["商品", "原价金额", "原价占比", "组合成交分摊"],
        [
            ("A", "60.00", "60%", "48.00"),
            ("B", "30.00", "30%", "24.00"),
            ("C", "10.00", "10%", "8.00"),
            ("合计", "100.00", "100%", "80.00"),
        ],
        [1500, 2340, 2340, 3180],
        center_columns={0, 1, 2, 3},
    )

    add_heading(doc, "5.6 营销互斥", 2)
    add_table(
        doc,
        ["权益/活动", "组合购交易处理"],
        [
            ("会员价", "不比较、不择优、不使用；组合商品按普通原价作为分摊基准。"),
            ("优惠券", "订单预览页不展示可用券，不锁券、不核销。"),
            ("折扣/限时折扣", "不叠加，不以折扣价作为组合购原价或成交价。"),
            ("满减/满赠/换购", "组合商品金额不计入门槛，也不享受优惠或赠品。"),
            ("积分抵现及其他营销", "不参与组合购订单的商品金额计算。"),
            ("运费/配送服务费", "不属于商品营销优惠，按商城普通履约规则单独计算。"),
        ],
        [2700, 6660],
    )
    add_body(doc, "消费者若希望使用其他价格权益，必须退出组合购流程并按普通商品路径购买。系统不得自动拆分组合后应用其他优惠。")
    add_callout(doc, "配置冲突", "默认按 SKU 维度校验：同一 SKU 只要存在任一未开始或进行中的组合购/其他促销即冲突，不因时间不重叠而放行；自然结束、手动终止或系统自动结束后释放占用。", fill=CAUTION, title_color="7A5A00")

    add_heading(doc, "5.7 订单预览", 2)
    add_table(
        doc,
        ["项目", "要求"],
        [
            ("请求参数", "activityId、activityVersion、setCount、最终履约 storeId、客户端商品集合摘要、配送方式、地址或自提门店。"),
            ("生命周期", "必须处于进行中且 terminatedAt/autoEndedAt 均为空；未开始和任一结束状态均拒绝预览。"),
            ("服务端重算", "按最终履约门店剔除不可交易 SKU；当前门店剩余少于 2 个时返回该店不可用，只有全局活动资格少于 2 个才原子结束活动；否则读取实时售价和库存并分摊。"),
            ("页面展示", "组合标识、活动名称、实际商品、固定组合价、配送费和应付金额；仅在原价合计高于组合价时展示原价与节省金额。"),
            ("数量锁定", "商品行不显示独立加减器；只允许返回组合面板调整套数。"),
            ("营销入口", "隐藏或禁用会员价、优惠券、满减、折扣和其他营销入口，并说明“组合购不与其他优惠同享”。"),
            ("配送变化", "地址、定位、手动门店、自提门店或配送方式导致履约门店变化后，废弃旧预览并重新计价。"),
            ("预览令牌", "预览成功返回一次性 previewToken、活动版本、实际商品集合哈希、门店价格快照和有效期。"),
        ],
        [2100, 7260],
    )

    add_heading(doc, "5.8 订单提交与库存", 2)
    for text in (
        "提交携带 previewToken 和客户端幂等键；同一幂等键只能生成一个订单。",
        "服务端再次校验活动进行中且无手动终止/自动结束标记；结束后已有 previewToken 立即失效，不允许提交新订单。",
        "若最终履约门店、活动版本、实际商品集合或售价快照与预览不一致，则令牌失效并要求重新预览，不能静默增删商品。",
        "库存按预览中的实际 SKU 数量一次性锁定；任一商品锁定失败则本次订单失败，不允许部分成单。",
        "提交成功后订单保存活动、实际商品集合、门店实时原价、固定组合价格、分摊、套数和门店快照；活动后续变化不回写历史订单。",
        "一次组合购订单只包含一种组合活动，不能混入普通商品或另一个组合。",
        "系统不校验单笔套数、单用户每日套数或活动周期累计套数；不得误用默认营销限购配置拦截订单。",
        "手动终止或系统自动结束只阻止新预览和新订单提交；结束前已创建的待支付订单仍可支付至订单自身超时，不自动取消。",
    ):
        add_list_item(doc, text, bullets)

    add_heading(doc, "5.9 订单详情", 2)
    add_table(
        doc,
        ["区域", "展示内容"],
        [
            ("订单标识", "明显展示“组合购”标签、活动名称和活动单号。"),
            ("价格摘要", "订单原价合计、组合成交总额、价差、运费/服务费和实付金额；价差不为正时不标记为优惠。"),
            ("商品明细", "实际成交商品图片、名称、规格、数量、履约门店原价和分摊成交金额；商品后续删除仍使用订单快照完整展示。"),
            ("组合信息", "购买套数、每套固定组合价格和含组合购订单只支持整单退款的说明。"),
            ("售后入口", "仅提供订单级“申请整单退款”，不在商品行提供独立退款入口。"),
        ],
        [2100, 7260],
    )

    add_heading(doc, "5.10 整单退款与异常售后", 2)
    add_callout(doc, "退款边界", "同一订单只要存在组合购商品，系统退款就必须覆盖整笔订单的全部商品和全部数量；无论组合购商品、普通商品、数量或套数都不能部分退款。", fill=ERROR_FILL, title_color="9B1C1C")
    add_table(
        doc,
        ["环节", "规则"],
        [
            ("申请入口", "订单级“申请整单退款”；组合商品行和普通商品行均不显示独立退款操作。"),
            ("申请范围", "自动锁定订单内全部商品、全部数量和全部套数，用户与客服均不可通过普通退款流程缩小范围。"),
            ("商品退款金额", "组合商品按成交分摊金额退款；订单内其他商品按各自成交快照退款，合计等于该订单应退商品金额。"),
            ("费用归集", "商品退款、原始运费、服务费及退回/拦截费用均由售后主单统一计算；商品分摊结果不因费用责任变化而改变。"),
            ("退款校验", "退款服务先识别订单级 containsCombination 标识，再核对订单快照与全部商品行；缺少任一行则拒绝创建退款单。"),
            ("异常售后", "缺货、漏发、破损、召回、配送丢失等单品异常仍不得发起部分退款；优先进入补发、换发、召回处置或客服补偿等独立流程。"),
            ("系统退款", "异常确需通过退款系统退回款项时仍按整单退款执行，不允许客服、仓库或定时任务绕过订单级限制。"),
            ("独立补偿", "客服补偿不生成商品退款单、不改写原订单商品分摊，并记录关联工单、原因、金额/权益和操作人。"),
            ("售后单结构", "创建一个订单级售后主单，并按仓库/包裹生成履约子任务；主单锁定整笔订单退款范围，子任务只管理拦截、退回和验收状态。"),
            ("未发货", "取消全部仓库出库任务并释放库存；所有子任务取消成功后按整单取消规则退款。"),
            ("部分发货/在途", "未发货包裹取消，在途包裹发起拦截；拦截失败转退货子任务。不得按已拦截包裹先行部分退款。"),
            ("部分签收", "已签收商品按整单范围生成退回任务，在途包裹继续拦截；免退货或召回例外仅改变履约动作，不缩小退款金额范围。"),
            ("跨仓完成", "各仓独立执行拦截、收货和质检；全部必需子任务达到可退款终态后，由售后主单一次性触发整单退款。"),
            ("费用执行", "履约子任务只回传状态和实际费用，不得自行计算应退金额或先退部分费用；售后主单在全部必需子任务闭环后一次性退回商品款及应退费用。"),
        ],
        [2100, 7260],
    )

    add_body(doc, "费用责任以整单售后主单的主责任类型为准，不因跨仓或拆包向消费者重复收费；同一订单存在多种原因时，混合责任按对消费者更有利的商家、平台或物流责任口径处理。")
    add_table(
        doc,
        ["责任/履约场景", "原始运费与服务费", "拦截与退回费用", "结算规则"],
        [
            ("未发货且未产生履约服务", "全部退还。", "无；已发生的取消处理费不得向消费者收取。", "全部出库任务取消成功后，由主单随商品款一次性退款。"),
            ("商家、平台或物流责任", "全部退还，包括缺货、漏发、破损、召回和配送丢失等场景。", "由商家、平台或物流责任方承担，消费者不垫付；已垫付的凭有效凭证补退。", "全部必需子任务闭环后，由主单一次性退还商品款、原始运费、服务费及消费者垫付款。"),
            ("用户责任且已开始履约", "原始运费和已实际发生的服务费不退；未发生部分应退。", "消费者最多承担一次标准退回运费；跨仓、拆包产生的新增退回费用由平台承担。", "主单统一核算应退商品款和费用，不允许各仓重复扣费。"),
            ("混合责任", "按商家、平台或物流责任口径全部退还。", "由商家、平台或物流责任方承担。", "责任口径在主单层级只确定一次，所有仓库/包裹子任务继承。"),
        ],
        [1900, 2500, 2600, 2360],
    )

    add_heading(doc, "5.11 状态与异常提示", 2)
    add_table(
        doc,
        ["场景", "处理", "推荐提示"],
        [
            ("活动未开始", "不展示入口、角标或活动标识；直接访问时按不可用处理。", "活动暂未开始"),
            ("活动终止/自动结束", "立即关闭购买按钮，使旧 previewToken 失效并重新查询可用组合。", "该组合购活动已失效"),
            ("当前门店不支持", "保留普通购买入口，组合购不可提交。", "当前门店暂不支持该组合购"),
            ("剔除后仍有至少 2 个 SKU", "更新组合面板和下一次预览，组合价格不变，旧令牌失效。", "组合商品已更新，请重新确认"),
            ("当前门店少于 2 个可交易 SKU", "仅隐藏当前门店入口并阻止预览/提交，不结束全门店活动；库存恢复且活动有效时可恢复展示。", "当前门店暂不支持该组合购"),
            ("全局活动资格少于 2 个 SKU", "永久自动结束活动、隐藏全部门店入口、释放促销占用并阻止预览/提交。", "该组合购活动已结束"),
            ("原价分摊基数为 0", "活动保持进行中，但当前门店阻止预览并触发运营告警。", "当前组合暂无法结算，请稍后再试"),
            ("多套库存不足", "保留实际商品集合，要求减少套数，旧 previewToken 失效。", "库存不足，请调整购买套数"),
            ("履约门店/原价变化", "废弃旧预览，按新门店售价重新分摊并要求确认；固定组合价格不变。", "商品价格已更新，请重新确认"),
            ("预览失败", "保留当前组合和套数，允许重试。", "订单金额计算失败，请稍后重试"),
            ("重复提交", "返回首个成功订单，不重复扣减库存。", "订单已提交，请勿重复操作"),
            ("尝试部分退款", "只要订单含组合购就拒绝部分退款流程。", "含组合购订单仅支持整单退款"),
        ],
        [2100, 4050, 3210],
    )

    add_heading(doc, "6. 数据与接口契约", 1)
    add_heading(doc, "6.1 交易快照", 2)
    add_table(
        doc,
        ["字段", "说明"],
        [
            ("activityId/activityNo/version", "商城组合购活动唯一标识、活动单号和成交版本，与 ERP 无关。"),
            ("storeId", "成交门店快照；活动范围虽为全部门店，订单仍记录实际履约门店。"),
            ("bundlePrice/setCount", "每套固定组合价格和购买套数。"),
            ("originalTotal/bundleTotal/differenceAmount", "实际商品门店原价合计、组合成交总额和价差；价差允许为 0 或负数。"),
            ("configuredProductIds/actualProductSetHash", "活动配置商品标识和本次实际成交商品集合摘要，用于审计动态剔除。"),
            ("productId/skuId/quantity", "每个实际成交 SKU 和实际购买数量。"),
            ("productName/specification/imageUrl", "成交时商品展示快照；商品下架或删除后订单与售后仍可完整展示。"),
            ("originalUnitPrice/originalLineAmount", "预览时最终履约门店售价和商品行原价金额快照。"),
            ("allocatedLineAmount", "按原价占比分摊后的商品行成交金额。"),
            ("allocationRuleVersion", "分摊算法版本，用于订单、退款和财务重算核对。"),
            ("containsCombination", "订单级布尔标识；为 true 时退款服务强制整单范围。"),
        ],
        [3150, 6210],
    )
    add_heading(doc, "6.2 逻辑接口", 2)
    add_table(
        doc,
        ["接口", "用途", "关键要求"],
        [
            ("按商品查询可用组合", "商品详情入口与组合列表", "只返回进行中、无结束标记且至少有 2 个可交易商品的商城活动。"),
            ("组合购详情", "展示本次可交易商品", "返回活动版本、配置 SKU 摘要、动态剔除后的实际 SKU、每套数量、固定组合价格和门店实时售价。"),
            ("组合购订单预览", "权威计价与资格校验", "以最终履约门店重算商品、原价、总价、分摊和库存；当前门店少于 2 个 SKU 时返回该店不可用，全局活动资格少于 2 个时原子结束活动。"),
            ("订单提交", "创建组合购订单", "使用 previewToken + 幂等键；再次校验无结束标记、履约门店和商品集合哈希。"),
            ("自动结束活动", "处理可售 SKU 不足", "原子写 autoEndedAt/reason、释放促销占用并清理缓存；接口幂等且活动不可恢复。"),
            ("订单详情", "展示成交快照", "返回组合标识、套数、价格与商品行分摊。"),
            ("整单退款", "创建含组合购订单退款单", "containsCombination=true 时只接受完整订单范围，拒绝任意商品行或部分数量参数。"),
            ("异常售后工单", "处理单品异常", "支持补发、召回处置和客服补偿；不得绕过退款服务创建部分退款。"),
            ("整单售后主单", "统筹部分履约/跨仓", "按仓库/包裹拆履约子任务，全部必需子任务闭环后一次性触发整单退款。"),
        ],
        [2400, 2700, 4260],
    )
    add_heading(doc, "6.3 服务端强校验", 2)
    for text in (
        "查询、预览、提交和退款均不得调用 ERP 组合价接口或依赖 ERP 状态。",
        "查询只返回进行中且 terminatedAt/autoEndedAt 均为空的活动；预览与提交必须再次验证。",
        "预览与提交必须验证 activityId + version + storeId + actualProductSetHash，禁止仅相信客户端价格和商品清单。",
        "实际商品集合由服务端按最终履约门店计算；当前门店少于 2 个 SKU 时只返回该店不可用，全局活动资格少于 2 个时以 CAS/事务原子结束活动，避免并发继续下单。",
        "履约门店实时售价用于原价、订单快照和分摊；原价合计不高于组合价时客户端只能展示固定组合价。",
        "分摊在服务端完成并保存算法版本；客户端只展示结果，不自行决定尾差。",
        "退款服务必须先读取 containsCombination，再读取订单成交快照；跨仓子任务只驱动履约状态，退款资金只能由整单主单一次触发。",
    ):
        add_list_item(doc, text, bullets)

    add_heading(doc, "7. 埋点与指标", 1)
    add_table(
        doc,
        ["事件", "关键属性", "用途"],
        [
            ("combination_exposure", "productId、skuId、storeId、activityCount", "评估入口覆盖。"),
            ("combination_click", "activityId、入口位置", "评估入口点击率。"),
            ("combination_switch", "fromActivityId、toActivityId", "分析多组合选择。"),
            ("combination_buy_click", "activityId、setCount、bundleTotal", "分析预览转化。"),
            ("combination_preview_result", "activityId、success、errorCode", "定位价格/库存失败。"),
            ("combination_order_success", "activityId、orderId、bundleTotal", "统计成交额与转化率。"),
            ("combination_refund_apply", "orderId、reason、fullOrder", "监控整单退款率和范围拦截。"),
            ("combination_product_excluded", "activityId、skuId、storeId、reason", "监控下架/停售/删除/库存不足导致的动态剔除。"),
            ("combination_auto_ended", "activityId、remainingEligibleSkuCount、reason", "监控全局活动资格 SKU 少于 2 个导致的系统自动结束。"),
            ("combination_aftersale_compensation", "orderId、skuId、reason、caseType", "监控补发与客服补偿。"),
        ],
        [2700, 4050, 2610],
    )
    add_body(doc, "核心指标包括：入口曝光率、预览/下单转化率、动态剔除率、自动结束量、手动终止拦截量、整单退款率、跨仓售后时长和异常售后率。")

    add_heading(doc, "8. 非功能要求", 1)
    add_table(
        doc,
        ["类别", "要求"],
        [
            ("一致性", "商品详情仅展示参考金额；预览、支付、订单和整单退款均使用服务端成交快照。"),
            ("幂等", "订单提交和整单退款分别支持幂等键，重复请求不重复下单或退款。"),
            ("性能", "可用组合查询 P95 不高于 1.5 秒；订单预览 P95 不高于 2 秒。"),
            ("可用性", "组合购服务异常不得阻断商品普通购买；不可回退到未经校验的缓存价格成交。"),
            ("时间", "活动资格由服务端中国标准时间判断，客户端时间不能决定是否成交。"),
            ("失效时效", "手动终止或系统自动结束后立即清理入口缓存或使用短 TTL；预览和提交必须实时拒绝。"),
            ("审计", "订单记录活动版本、门店、配置/实际商品集合、分摊算法版本、预览令牌关联和退款范围。"),
        ],
        [1800, 7560],
    )

    add_heading(doc, "9. 验收标准", 1)
    add_table(
        doc,
        ["编号", "验收项", "通过标准"],
        [
            ("AC-01", "ERP 独立", "商城查询、预览、提交和退款均不读取 ERP 组合价数据。"),
            ("AC-02", "生命周期", "未开始不展示；手动终止或系统自动结束后立即隐藏，旧预览令牌不能提交新订单。"),
            ("AC-03", "履约门店", "地址、定位或门店变化后以最终履约门店重新查询和计价，并使旧预览令牌失效。"),
            ("AC-04", "系统组单", "组合面板无消费者勾选、删除、替换或主商品逻辑；实际商品集合由服务端计算。"),
            ("AC-05", "加入购物车", "首次点击加入 1 套完整组合；连续点击 N 次后同一活动版本与履约门店的组合套数和购物车角标均增加 N，商品明细仍保持整组绑定。"),
            ("AC-06", "原价下降", "原价合计等于或低于组合价时活动和价格不变，页面只展示固定组合价，无优惠、节省或原价划线。"),
            ("AC-07", "营销互斥", "会员价、优惠券、折扣、满减和其他营销均不参与展示、计算、锁定或核销。"),
            ("AC-08", "活动冲突", "商品存在任一未开始/进行中的组合购或其他促销即冲突，不因时间不重叠放行；结束后释放占用。"),
            ("AC-09", "动态商品", "不可售 SKU 被剔除且组合价不变；当前门店少于 2 个可交易 SKU 时仅隐藏该店入口，全局活动资格少于 2 个时活动永久自动结束并释放占用。"),
            ("AC-10", "金额分摊", "分摊按商品原价金额占比计算，按分处理尾差，分摊合计严格等于组合成交总额。"),
            ("AC-11", "订单快照", "订单可追溯活动、履约门店、商品名称/规格/图片、原价和分摊；商品删除后仍完整展示。"),
            ("AC-12", "整单退款", "只要订单含组合购，用户和客服均无法选择部分商品、数量或套数，退款范围固定为整笔订单。"),
            ("AC-13", "异常售后", "缺货、漏发、破损、召回或配送丢失通过补发/客服补偿工单处理；系统退款仍强制整单。"),
            ("AC-14", "不限套数", "无组合购单笔、单用户每日或活动周期累计套数限制；只受实时库存与普通合规规则约束。"),
            ("AC-15", "预览一致", "门店、活动状态/版本、实际商品集合或售价变化时旧令牌失效，重新确认后方可提交。"),
            ("AC-16", "服务端权威", "篡改客户端商品集合、数量、门店或价格不能影响服务端预览和成交结果。"),
            ("AC-17", "库存原子性", "预览中的任一实际商品锁库失败时订单失败，不产生部分订单或部分库存扣减。"),
            ("AC-18", "幂等", "重复提交和重复整单退款请求分别只产生一个订单和一个退款单。"),
            ("AC-19", "零分摊基数", "实际商品原价合计为 0 时不执行除法或均分，当前门店预览失败并告警，活动状态不被改写。"),
            ("AC-20", "终止前订单", "活动终止前已创建的订单不被自动取消；待支付订单沿用订单自身的支付超时规则。"),
            ("AC-21", "跨仓售后", "部分发货、部分签收或跨仓时仅生成一个售后主单；全部必需子任务闭环后一次性整单退款。"),
            ("AC-22", "履约拦截", "未发货取消、在途拦截、已签收退回分别形成子任务，任何包裹或仓库均不能先行部分退款。"),
            ("AC-23", "费用口径", "未履约费用全退；商家、平台或物流责任时原始费用全退且责任方承担退回费用；用户责任且已履约时已发生费用不退且消费者最多承担一次标准退回运费；混合责任按商责口径，所有金额只由主单计算。"),
        ],
        [1050, 2100, 6210],
        center_columns={0},
    )

    add_heading(doc, "10. 技术实施注意事项", 1)
    for text in (
        "商城组合购的 activityId 必须是商城域内部标识，不复用 ERP wareid、warecode 或 collocationId 映射关系。",
        "门店普通售价必须来自统一权威价格服务；浏览可用选择/LBS 门店，但结算统一切换为地址或自提方式确定的最终履约门店。",
        "可交易商品集合计算由查询、详情、预览和提交复用；必须区分单店库存不足与全局商品资格失效，只有后者少于 2 个时以原子状态迁移写 autoEndedAt。",
        "手动终止和自动结束分别通过 terminatedAt/autoEndedAt 覆盖时间状态，缓存失效与交易实时校验并行。",
        "分摊尾差算法及版本须由订单、支付、发票、财务和退款服务共同实现或复用同一服务。",
        "订单域保存 containsCombination 和不可变商品展示/价格快照；售后主单与跨仓子任务采用状态机，资金退款只由主单触发。",
        "处方、禁限售、配送等普通合规规则仍为强校验；不合规 SKU 的剔除或整单拦截策略按法规级别执行，不得因组合购弱化。",
    ):
        add_list_item(doc, text, bullets)

    doc.core_properties.title = "千金健康商城端组合购PRD"
    doc.core_properties.subject = "商城组合购展示、交易、分摊与整单退款"
    doc.core_properties.author = "千金健康商城产品团队"
    doc.core_properties.keywords = "组合购, 商城, 固定总价, 门店售价, 金额分摊, 整单退款, PRD"
    doc.save(MALL_OUT)


def main():
    build_admin_prd()
    build_mall_prd()
    print(ADMIN_OUT)
    print(MALL_OUT)


if __name__ == "__main__":
    main()
