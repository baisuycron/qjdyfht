from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

import build_combination_price_reference_prd as base


ROOT = Path(".")
ASSET_DIR = ROOT / "requirements_assets" / "combo_v23"
CROP_DIR = ASSET_DIR / "cropped"
BASE_OUTPUT = ROOT / "combination_price_reference_v22.docx"
OUTPUT = ROOT / "combination_price_fusion_screenshot_v23.docx"


def iter_table_paragraphs(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                yield paragraph
            for nested in cell.tables:
                yield from iter_table_paragraphs(nested)


def iter_all_paragraphs(doc):
    yield from doc.paragraphs
    for table in doc.tables:
        yield from iter_table_paragraphs(table)
    for section in doc.sections:
        yield from section.header.paragraphs
        yield from section.footer.paragraphs


def replace_in_runs(paragraph, old, new):
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)


def set_paragraph_text(paragraph, text):
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def apply_copy_and_rule_updates(doc):
    paragraphs = list(iter_all_paragraphs(doc))
    for paragraph in paragraphs:
        replace_in_runs(paragraph, "组合购", "组合价")
        replace_in_runs(paragraph, "V2.2", "V2.3")
        replace_in_runs(paragraph, "2026-08-26", "2026-08-28")

    replacements = {
        "千金小程序需求-组合价": "千金小程序需求-组合价（后台与小程序融合版）",
        "后台组合价 · 商城组合价 · 小程序交易链路": "后台配置 · 小程序展示 · 购物车、订单与售后链路",
        "数据关系修订版（含待确认问题）": "后台与小程序融合截图版（含待确认问题）",
        "千金大药房后台项目已经具备“组合价”入口、活动列表、查询、新增、编辑、查看和手动结束等原型功能；千金大药房商城项目已经具备商品详情优惠组合、组合商品列表、加入购物车、立即购买、购物车整组展示、确认订单和商品组合单价展示等原型能力。": "千金大药房后台项目已经具备“组合价”入口、活动列表、查询、新增、编辑、查看和手动结束等原型功能；小程序原型已经具备商品详情组合价入口、优惠组合、组合商品列表、加入购物车、立即购买、购物车整组展示和确认订单等原型能力。本文将后台配置与小程序消费链路融合描述。",
        "保证商品详情、购物车、确认订单、提交订单、订单详情和售后使用同一活动版本、履约门店、商品组合单价和成交金额口径。": "保证商品详情、购物车、确认订单、提交订单、订单详情和售后使用同一活动版本、履约门店、固定组合价与商品行成交金额口径。",
        "根据后台同一活动和最终履约门店，重新检查实际商品清单、门店原价、活动组合价格、商品组合单价、库存和商品行成交金额。": "根据后台同一活动和最终履约门店，重新检查实际商品清单、门店实时原价、固定组合价和库存，并按实际成交商品门店原价金额占比计算商品行成交金额。",
        "展示组合标识、套数、后台商品组合单价和成交金额；含组合价订单只允许整单退款。": "展示组合价标识、套数、固定组合价和商品行成交金额；只要订单中含组合价商品，售后范围固定为整笔订单。",
        "商城直接使用后台活动中的组合价格、商品组合单价和每组数量；最终履约门店实时普通售价仅用于展示原价和计算优惠，原价变化不修改后台组合价格。": "商城直接使用后台活动中的固定组合价格、配置商品和每组数量；最终履约门店实时普通售价用于展示原价、判断可售性并按比例分配商品行成交金额，原价变化不修改后台固定组合价格。",
        "订单保存后台活动信息、履约门店、实际商品清单、门店原价、组合价格和商品行成交金额；含组合价订单仅支持整单退款。": "订单保存后台活动信息、履约门店、实际商品清单、门店原价、固定组合价格、分配比例和商品行成交金额；含组合价订单仅支持整单退款。",
        "活动名称、活动编号、活动时间、活动状态、组合价格、组合商品、商品组合单价及每组数量。": "活动名称、活动编号、活动时间、活动状态、固定组合价格、组合商品及每组数量；商品行展示金额由交易侧按当前履约门店原价比例计算。",
        "直接使用后台活动中的组合价格、商品组合单价和数量，商城端不得重新配置或改写。": "直接使用后台活动中的固定组合价格、商品和数量；交易侧只能计算本次商品行成交金额，不得回写或改写后台活动配置。",
        "活动名称、商品数量、组合价格、商品图片/名称/规格、商品组合单价、每组数量。": "活动名称、商品数量、固定组合价格、商品图片/名称/规格、商品行展示金额和每组数量。",
        "组合价格、套数、组合商品总件数、后台商品组合单价和商品行成交金额。": "固定组合价格、套数、组合商品总件数和本次商品行成交金额。",
        "地址、定位、配送方式或门店变化使旧预览失效；最终履约门店少于 2 个可交易商品时隐藏入口。": "地址、定位、配送方式或门店变化使旧预览失效；不可售商品在当前门店隐藏，剩余商品不少于 2 个时固定组合价保持不变并重新分配商品行成交金额，少于 2 个时隐藏该门店入口。",
        "商城组合价来自后台组合价活动。活动编号、名称、时间、状态、组合商品、商品组合单价、每组数量和组合价格必须保持一致；商城端只负责展示和交易，不能另建活动或修改后台配置。": "小程序组合价来自后台组合价活动。活动编号、名称、时间、状态、配置商品、每组数量和固定组合价格必须保持一致；小程序只负责展示和交易，不能另建活动或修改后台配置。本次商品行成交金额属于订单计算结果，不回写活动。",
        "从组合商品列表进入确认订单，并带上后台组合价活动编号、活动版本、履约门店和购买套数。": "从组合商品列表进入确认订单，并带上后台组合价活动编号、活动版本、履约门店、固定组合价和购买套数。",
        "系统读取后台同一组合价活动，重新检查实际商品清单、最终履约门店普通售价、活动组合价格、商品组合单价、每组数量和库存，并返回本次订单预览凭证。": "系统读取后台同一组合价活动，重新检查实际商品清单、最终履约门店普通售价、固定组合价、每组数量和库存，按实际成交商品门店原价金额占比计算商品行成交金额，并返回本次订单预览凭证。",
        "保存后台活动编号、活动名称、活动版本、履约门店、配置商品和实际成交商品、商品组合单价、每组数量、购买套数、门店原价、组合价格和商品行成交金额。": "保存后台活动编号、活动名称、活动版本、履约门店、配置商品和实际成交商品、每组数量、购买套数、门店原价、固定组合价格、分配比例、尾差归集商品和商品行成交金额。",
        "列表和详情展示“组合价”标识；详情展示每套组合价、购买套数、实际商品与商品行成交金额。": "列表和详情展示“组合价”标识；详情展示每套固定组合价、购买套数、实际成交商品、门店原价、分配比例与商品行成交金额。",
        "报表按后台商品组合单价、每组数量和购买套数统计商品成交金额；所有商品行合计必须等于组合价成交总额。": "报表以订单快照中的商品行成交金额统计；商品行按实际成交商品门店原价金额占比分配，所有商品行合计必须等于组合价成交总额。",
        "商城直接使用后台同一活动中保存的组合价格、商品组合单价和每组数量，商城端不能另设一套价格。": "商城直接使用后台同一活动中保存的固定组合价格、配置商品和每组数量；后台商品组合单价仅用于活动配置与固定组合价格汇总，商城端不能另设一套活动价格。",
        "每个商品行按“后台配置的商品组合单价 × 每组数量 × 购买套数”计算。所有商品行相加必须等于组合成交总金额。": "按实际成交商品的“最终履约门店售价 × 每组数量 × 购买套数”计算原价金额占比，再用该占比分配组合成交总额；金额保留到分，所有商品行合计必须等于组合成交总额。",
        "后台保存活动前先检查商品行金额合计与活动组合价格是否一致；商城预览和提交订单时再次检查，不一致时停止下单并刷新后台活动数据。": "后台保存活动前检查配置商品组合单价汇总与固定组合价格一致；商城预览和提交订单时检查分配后的商品行成交金额合计严格等于组合成交总额，不一致时停止下单并重新计算。",
        "后台活动版本、组合价格、商品清单、商品组合单价、门店或售价发生变化时，原来的订单预览结果失效；系统读取最新后台活动并提示用户重新确认。": "后台活动版本、固定组合价格、商品清单、最终履约门店或门店售价发生变化时，原订单预览结果失效；系统重新读取活动、判断可售商品并计算商品行成交金额后提示用户确认。",
        "后台活动中 A、B、C 三个商品每组各 1 件，商品组合单价分别为 48 元、24 元和 8 元，活动组合价格自动汇总为 80 元。购买 2 套时，三个商品行成交金额分别为 96 元、48 元和 16 元，合计 160 元，与“80 元 × 2 套”一致。": "后台固定组合价格为 80 元。最终履约门店中 A、B、C 每套原价金额分别为 60 元、30 元和 10 元，占比为 60%、30% 和 10%，则每套商品行成交金额为 48 元、24 元和 8 元；购买 2 套时为 96 元、48 元和 16 元，合计 160 元。金额保留到分后产生的尾差按 5.9 待确认规则归集。",
        "仅某门店因库存不足导致可交易商品少于 2 个时，只隐藏该门店入口；库存恢复且活动仍有效时可恢复。": "当前门店中个别商品因库存、下架、停售、删除或交易限制不可售时隐藏该商品；剩余可售商品不少于 2 个时固定组合价保持不变并按剩余商品原价占比重新分配，少于 2 个时只隐藏该门店入口。库存恢复且活动仍有效时可恢复展示。",
        "某门店的部分组合商品因库存、下架、停售、禁限售、处方或配送规则不能交易，但剩余商品仍不少于 2 个时，是去掉不可交易商品继续购买，还是整组不允许购买？如继续购买，活动组合价格及被去掉商品的组合单价怎么处理？": "商品行成交金额按门店原价占比分配并保留到分后，尾差应固定归集到哪一商品行（如最高金额商品、最后一行或商品编码最小行）？",
        "直接影响商城展示、商品行成交金额、活动组合价格、订单和对账是否一致。": "需要确定可重复计算、可审计的唯一规则，保证订单、退款、发票和财务对账一致。",
        "P-01、P-02、P-03 和 P-04 会直接改变商城购物车、订单、商品数量或成交金额计算方式，应在研发排期前确认；在确认前不得把当前演示数据、旧版小程序流程或后台原型字段描述成已确定的正式规则。": "P-01、P-02、P-03 和 P-04 会直接改变购物车、订单、商品数量或商品行成交金额计算方式，应在研发排期前确认；在确认前不得把当前演示数据、旧版流程或原型字段描述成已确定的正式接口规则。",
        "每个商品行按后台商品组合单价、每组数量和购买套数计算，所有商品行合计严格等于组合成交总额。": "商品行按实际成交商品门店原价金额占比分配组合成交总额，按确定的尾差规则归集后，所有商品行合计严格等于组合成交总额。",
        "地址、定位或门店变化后重新查询后台活动、商品、售价和库存，并重新计算商品行成交金额，原订单预览结果失效。": "地址、定位或门店变化后重新查询后台活动、商品、售价和库存，隐藏不可售商品并重新按门店原价占比分配商品行成交金额，原订单预览结果失效。",
    }

    for paragraph in paragraphs:
        text = paragraph.text
        if text in replacements:
            set_paragraph_text(paragraph, replacements[text])


def crop_mobile_screenshots():
    CROP_DIR.mkdir(parents=True, exist_ok=True)
    sources = [
        "05_小程序商品详情组合价入口.png",
        "06_小程序组合商品列表.png",
        "08_小程序组合价购物车.png",
    ]
    outputs = []
    for name in sources:
        source = ASSET_DIR / name
        target = CROP_DIR / name
        with Image.open(source) as image:
            image.crop((470, 18, 810, 712)).save(target, optimize=True)
        outputs.append(target)
    return outputs


def set_figure_alt(inline_shape, alt_text):
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", alt_text)
    doc_pr.set("title", alt_text)


def add_figure_elements(doc, image_path, width, caption, alt_text):
    image_paragraph = doc.add_paragraph()
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.paragraph_format.space_before = Pt(4)
    image_paragraph.paragraph_format.space_after = Pt(4)
    image_paragraph.paragraph_format.keep_with_next = True
    shape = image_paragraph.add_run().add_picture(str(image_path), width=Inches(width))
    set_figure_alt(shape, alt_text)

    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.space_before = Pt(0)
    caption_paragraph.paragraph_format.space_after = Pt(2)
    caption_paragraph.paragraph_format.keep_with_next = True
    base.set_run_font(caption_paragraph.add_run(caption), size=9.2, bold=True, color=base.DARK_BLUE)

    note_paragraph = doc.add_paragraph()
    note_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note_paragraph.paragraph_format.space_before = Pt(0)
    note_paragraph.paragraph_format.space_after = Pt(8)
    base.set_run_font(
        note_paragraph.add_run("原型截图，采集于 2026-08-28；界面数据为模拟数据，不代表正式生产接口已接入。"),
        size=8.5,
        color=base.GRAY,
        italic=True,
    )
    return [image_paragraph._p, caption_paragraph._p, note_paragraph._p]


def insert_figures_after_heading(doc, heading_text, figures):
    target = next((paragraph for paragraph in doc.paragraphs if paragraph.text == heading_text), None)
    if target is None:
        raise RuntimeError(f"heading not found: {heading_text}")
    reference = target._p
    for figure in figures:
        elements = add_figure_elements(doc, **figure)
        for element in elements:
            reference.addnext(element)
            reference = element


def add_prototype_note(doc):
    target = next((paragraph for paragraph in doc.paragraphs if paragraph.text == "一、需求背景"), None)
    if target is None:
        raise RuntimeError("background heading not found")
    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(4)
    note.paragraph_format.space_after = Pt(8)
    note.paragraph_format.left_indent = Inches(0.08)
    p_pr = note._p.get_or_add_pPr()
    shading = p_pr.find(qn("w:shd"))
    if shading is None:
        from docx.oxml import OxmlElement

        shading = OxmlElement("w:shd")
        p_pr.append(shading)
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), "FFF7E3")
    base.set_run_font(note.add_run("截图证据边界："), size=10, bold=True, color="7A5A00")
    base.set_run_font(
        note.add_run("本文截图来自当前后台与小程序本地原型，用于说明页面结构和交互方向；正式上线仍需补齐后台持久化、小程序读取同一活动、订单计价与售后服务接口。"),
        size=10,
        color="333333",
    )
    target._p.addnext(note._p)


def build():
    base.OUTPUT = BASE_OUTPUT
    base.build()
    doc = Document(BASE_OUTPUT)
    apply_copy_and_rule_updates(doc)
    cropped_detail, cropped_list, cropped_cart = crop_mobile_screenshots()
    add_prototype_note(doc)

    insert_figures_after_heading(
        doc,
        "5.2 组合价 - 活动列表",
        [
            {
                "image_path": ASSET_DIR / "01_后台组合价活动列表.png",
                "width": 6.2,
                "caption": "图 1  后台组合价活动列表",
                "alt_text": "后台组合价活动列表，包含筛选条件、活动状态和查看编辑结束操作",
            }
        ],
    )
    insert_figures_after_heading(
        doc,
        "5.3 组合价 - 活动列表 - 活动详情",
        [
            {
                "image_path": ASSET_DIR / "02_后台组合价活动详情.png",
                "width": 6.2,
                "caption": "图 2  后台组合价活动详情",
                "alt_text": "后台组合价活动详情，展示活动时间、固定组合价格和组合商品明细",
            },
            {
                "image_path": ASSET_DIR / "04_后台组合商品选择.png",
                "width": 6.2,
                "caption": "图 3  后台组合商品选择与促销冲突提示",
                "alt_text": "后台组合商品选择弹窗，展示商品筛选、促销冲突和最多九个商品限制",
            },
        ],
    )
    insert_figures_after_heading(
        doc,
        "5.4 小程序端 - 展示规则",
        [
            {
                "image_path": cropped_detail,
                "width": 3.0,
                "caption": "图 4  小程序商品详情组合价入口",
                "alt_text": "小程序商品详情页的组合价优惠组合卡片和去购买组合套装按钮",
            },
            {
                "image_path": cropped_list,
                "width": 3.0,
                "caption": "图 5  小程序组合商品列表",
                "alt_text": "小程序组合商品列表，展示固定组合价、三个商品和加入购物车立即购买操作",
            },
            {
                "image_path": cropped_cart,
                "width": 3.0,
                "caption": "图 6  小程序购物车中的组合价整组",
                "alt_text": "小程序购物车中的组合价活动分组、套数、优惠组合和立即支付区域",
            },
        ],
    )

    doc.core_properties.title = "千金小程序需求-组合价（后台与小程序融合版）"
    doc.core_properties.subject = "后台组合价与小程序组合价融合需求（含原型截图）"
    doc.core_properties.keywords = "组合价, 后台, 小程序, 同一活动, 固定组合价, 商品行金额分配, 整单退款, PRD"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
