from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parent.parent


def all_text(doc):
    parts = [paragraph.text for paragraph in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def validate(path, required, forbidden):
    doc = Document(path)
    text = all_text(doc)
    errors = []
    for phrase in required:
        if phrase not in text:
            errors.append(f"missing required phrase: {phrase}")
    for phrase in forbidden:
        if phrase in text:
            errors.append(f"forbidden legacy phrase present: {phrase}")
    if "\ufffd" in text:
        errors.append("replacement character detected")
    for marker in ("TODO", "TBD", "{{", "}}", ":codex-file-citation", "PLACEHOLDER"):
        if marker in text:
            errors.append(f"internal placeholder detected: {marker}")
    if any(len(table.rows) == 0 or len(table.columns) == 0 for table in doc.tables):
        errors.append("empty table detected")
    if any(not section.header.paragraphs or not section.footer.paragraphs for section in doc.sections):
        errors.append("section header/footer structure missing")
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    if not headings:
        errors.append("no real heading styles found")
    print(f"{path.name}: paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} headings={len(headings)} chars={len(text)}")
    if errors:
        for error in errors:
            print(f"  ERROR: {error}")
        return False
    print("  status=OK")
    return True


def main():
    admin_ok = validate(
        ROOT / "后台系统组合购PRD_V1.2.docx",
        required=(
            "商城组合购由商城后台独立创建和维护",
            "与 ERP 组合价无任何数据同步",
            "固定整组总价",
            "每个组合至少 2 个、最多 9 个",
            "适用于全部门店",
            "0 < 组合价格 < 商品原价合计",
            "不参与任何其他价格权益",
            "只允许整单退款",
            "活动单号",
            "活动起止时间",
            "商品名称",
            "商品编码",
            "新增组合购",
            "选择本页/取消本页",
            "storeScope=ALL",
            "未开始：查看、编辑、终止",
            "未开始和进行中",
            "活动名称和活动单号均建立唯一约束",
            "促销占用",
            "开始时间允许早于保存时刻",
            "结束时间必须晚于当前时间",
            "SELLABLE_SKU_LT_2",
            "数据库操作日志",
            "不判断活动时间是否重叠",
            "AC-23",
        ),
        forbidden=(
            "ERP 是活动创建",
            "ERP 为组合价活动的唯一维护入口",
            "主商品必须固定选中",
            "至少再选择1个商品",
            "新增成功后的活动不可通过本期后台修改",
            "结束时间晚于开始时间即可保存",
        ),
    )
    mall_ok = validate(
        ROOT / "商城端组合购PRD_V1.2.docx",
        required=(
            "商城组合购与 ERP 完全独立",
            "固定组合价",
            "动态剔除",
            "适用于全部门店",
            "0 < 组合价格 < 商品原价合计",
            "不参与任何其他价格权益",
            "只允许整单退款",
            "用户当前选择或 LBS 定位门店",
            "活动未开始时不展示入口、角标或任何活动标识",
            "组合价格保持不变",
            "少于 2 个时活动永久自动结束",
            "不设置组合购活动级上限",
            "默认按 SKU 维度校验",
            "不因时间不重叠而放行",
            "最终履约门店",
            "只展示固定组合价",
            "原价分摊基数为 0",
            "商品行金额分摊",
            "Aᵢ* = T × wᵢ",
            "按分截断",
            "分配尾差",
            "申请整单退款",
            "异常售后",
            "补发、换发、召回处置或客服补偿",
            "售后主单",
            "全部必需子任务",
            "商家、平台或物流责任",
            "消费者最多承担一次标准退回运费",
            "混合责任",
            "商品后续删除仍使用订单快照完整展示",
            "组合购不与其他优惠同享",
            "每点击一次加入购物车增加 1 套完整组合",
            "购物车角标",
            "购物车中的组合商品整组绑定",
            "AC-23",
        ),
        forbidden=(
            "主商品默认选中",
            "非主商品可勾选",
            "按更低金额生效",
            "支持部分退款",
            "若 SKU 失效，整个组合不可购买",
            "组合价格必须 >0 且低于当前普通原价合计，否则不能预览或提交",
            "至少保留 1 个商品时可成交",
            "剩余至少 1 个商品时可下单",
            "0 个时隐藏活动",
        ),
    )
    raise SystemExit(0 if admin_ok and mall_ok else 1)


if __name__ == "__main__":
    main()
