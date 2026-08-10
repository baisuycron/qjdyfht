from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parent.parent / "组合价需求文档.docx"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
GRAY = "666666"
LIGHT_BLUE = "E8EEF5"
CALLOUT = "F4F6F9"
CONTENT_WIDTH_DXA = 9360


def set_run_font(run, size=11, bold=False, color="000000", name="Microsoft YaHei"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths, indent=120):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    for tag, value in (("w:tblW", sum(widths)), ("w:tblInd", indent)):
        node = tbl_pr.first_child_found_in(tag)
        if node is None:
            node = OxmlElement(tag)
            tbl_pr.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    for i, width in enumerate(widths):
        table._tbl.tblGrid.gridCol_lst[i].set(qn("w:w"), str(width))
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[i]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table, color="D9E2F0"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def write_cell(cell, text, bold=False, color="000000", size=10.2, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.12
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, color=color)


def add_table(doc, headers, rows, widths, first_col_center=True):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table)
    for i, label in enumerate(headers):
        set_cell_shading(table.rows[0].cells[i], LIGHT_BLUE)
        write_cell(table.rows[0].cells[i], label, bold=True, color=INK, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.CENTER if first_col_center and i == 0 else WD_ALIGN_PARAGRAPH.LEFT
            write_cell(cells[i], value, align=align)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    return table


def add_body(doc, text, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    set_run_font(p.add_run(text), size=11)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    set_run_font(p.add_run(text), size=11)
    return p


def add_h1(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    set_run_font(p.add_run(text), size=16, bold=True, color=BLUE)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    set_run_font(p.add_run(text), size=13, bold=True, color=BLUE)
    return p


def add_callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    set_table_borders(table, "D5E1F0")
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.12
    set_run_font(p.add_run(f"{title}："), size=10.5, bold=True, color=DARK_BLUE)
    set_run_font(p.add_run(text), size=10.5, color="333333")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_meta_row(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.1
    set_run_font(p.add_run(f"{label}："), size=10.5, bold=True, color=INK)
    set_run_font(p.add_run(value), size=10.5, color="333333")


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, end))


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    for style_name, size, color in (("Normal", 11, "000000"), ("Heading 1", 16, BLUE), ("Heading 2", 13, BLUE)):
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    set_run_font(header.add_run("千金大药房 | 产品需求文档"), size=9, color=GRAY)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(footer.add_run("第 "), size=9, color=GRAY)
    add_page_field(footer)
    set_run_font(footer.add_run(" 页"), size=9, color=GRAY)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(10)
    title.paragraph_format.space_after = Pt(4)
    title.paragraph_format.keep_with_next = True
    set_run_font(title.add_run("组合价需求文档"), size=23, bold=True, color=INK)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    set_run_font(subtitle.add_run("覆盖 ERP 配置、管理后台查询与小程序商城展示/结算"), size=12, color=GRAY)
    add_meta_row(doc, "所属模块", "营销中心 > 门店促销 > 组合价")
    add_meta_row(doc, "版本", "V1.0")
    add_meta_row(doc, "日期", "2026-08-07")
    add_meta_row(doc, "维护边界", "活动与组合规则由 ERP 维护；后台负责查询与详情展示；小程序商城负责展示、凑单引导和结算呈现")
    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(8)
    ppr = rule._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), BLUE)
    borders.append(bottom)
    ppr.append(borders)

    add_h1(doc, "一、需求背景")
    add_body(doc, "为提升关联商品的连带购买，门店需要将指定商品按指定数量组成一个固定价格的销售组合。消费者在小程序商城凑齐组合商品后，可按组合价购买；活动、组合明细、适用门店、时段和适用对象需要与 ERP 配置保持一致。")
    add_body(doc, "组合价不在管理后台重复维护。ERP 是活动创建、编辑、启停、审批及商品范围维护的唯一来源；管理后台用于查询和核对，商城与结算端以已同步且当前有效的组合规则为准。")
    add_callout(doc, "核心原则", "同一组合必须同时满足全部指定商品及数量，才可享受组合固定价；商品展示、购物车、订单提交和售后关联的价格口径必须一致。")

    add_h1(doc, "二、需求目标")
    for item in (
        "让 ERP 能够集中维护组合价活动、组合商品、数量、固定组合价、门店、时间与适用对象，并将有效数据下发至商城。",
        "让运营与门店人员在管理后台快速查询活动状态、适用范围和组合商品明细，核对商城生效数据。",
        "让消费者在小程序商城清楚识别可参与组合价的商品，并得到凑齐组合所需商品的引导。",
        "保障组合价、会员价与其他促销的优先级判断一致，避免商品页、购物车和结算页出现价格不一致。",
        "确保订单记录可追溯组合活动与组合明细，为履约、退款和对账提供依据。",
    ):
        add_bullet(doc, item)

    add_h1(doc, "三、业务流程")
    add_body(doc, "组合价活动由 ERP 配置并同步至营销服务，管理后台提供只读查询与详情核对；小程序商城基于有效规则展示、凑单并在结算时校验组合资格。")
    add_table(doc, ["步骤", "处理方", "业务动作", "产出"], [
        ("1", "ERP", "创建或维护组合价活动，配置组合商品、每组数量、组合固定价、门店、时间、适用对象与叠加规则。", "可同步的组合价活动与组合明细"),
        ("2", "数据同步", "将活动主数据、组合商品明细及状态同步至营销服务；失败时保留最近一次成功数据并告警。", "商城可读取的有效组合规则"),
        ("3", "管理后台", "按条件查询已同步活动，查看基本规则、组合商品和状态；不提供活动维护。", "活动清单与只读详情"),
        ("4", "小程序商城", "在商品页、购物车识别命中组合的商品，展示组合价与凑单进度，引导消费者凑齐组合。", "组合价展示与凑单提示"),
        ("5", "结算/订单", "提交订单前重新校验门店、时间、商品、数量和库存；满足条件按组合价计价并保存活动关联。", "实际成交价与可追溯订单记录"),
    ], [700, 1300, 4400, 2960])
    add_callout(doc, "资格判断", "组合资格以“同一订单、同一履约门店、有效活动期间、满足全部指定商品及数量”为准；任一商品缺失、数量不足、无货或不再适用时，该组合不生效。")

    add_h1(doc, "四、产品方案概要")
    add_table(doc, ["范围", "方案说明"], [
        ("ERP", "提供组合价活动的创建、编辑、复制、启停、审批及商品组合维护；维护活动时间、门店、适用对象、组合固定价和叠加控制，并下发活动与组合明细。"),
        ("数据同步", "以 ERP 为唯一数据源同步活动主表、组合组、商品及数量明细、适用范围和状态；为商城提供可用的查询与价格计算数据。"),
        ("管理后台", "提供活动列表、条件查询、重置和只读详情；用于核验 ERP 同步结果，不提供新增、编辑、删除或回写 ERP 的操作。"),
        ("小程序商城", "在商品详情、购物车等承载交易价格的页面展示组合价、组合商品与凑单提示；完成组合时在购物车及结算页展示优惠结果。"),
        ("计价与订单", "以完整组合为计价单元；按 ERP 下发的组合固定价生效，并在订单中保存活动编号、组合编号、商品数量和分摊金额/规则。"),
    ], [2300, 7060])

    add_h1(doc, "五、产品方案细节")
    add_h2(doc, "5.1 ERP - 组合价活动维护")
    add_body(doc, "ERP 为组合价活动的唯一维护入口。活动生效前应完成组合完整性、商品可售性、时间和门店范围校验；审批、权限和启停沿用 ERP 现有机制。")
    add_table(doc, ["区域", "字段/动作", "规则"], [
        ("基本信息", "活动名称、活动编号、活动状态、活动时间、适用门店、适用对象、参与日期/星期", "活动编号全局唯一；状态按时间和 ERP 启停状态综合计算；门店、对象及日期条件均须满足。"),
        ("组合信息", "组合编号、组合名称、组合固定价、原价合计、限购规则、排序", "组合固定价为该完整组合的总成交价；原价合计用于展示优惠，不参与最终结算。"),
        ("组合商品", "商品编码、商品名称、规格、每组数量、销售价、商品状态", "每组商品与数量为必填；商品下架、禁售或不适用当前门店时，组合不得在商城生效。"),
        ("活动操作", "新建、编辑、复制、启用、停用、提前结束、审批", "仅 ERP 提供上述维护操作；变更后按同步机制下发，商城以最近一次成功同步且有效的数据为准。"),
        ("冲突校验", "商品促销冲突、组合重复、价格合法性", "同一商品同一时间仅参加一个促销活动；ERP 在保存/启用时拦截冲突。组合固定价必须大于 0 且低于组合原价合计。"),
    ], [1400, 3300, 4660])

    add_h2(doc, "5.2 管理后台 - 活动列表")
    add_body(doc, "入口路径：营销中心 > 门店促销 > 组合价。页面默认展示 ERP 已成功同步的组合价活动，仅支持查询、重置与查看详情。")
    add_table(doc, ["区域", "字段/动作", "规则"], [
        ("查询条件", "活动名称", "支持模糊查询。"),
        ("查询条件", "活动编号", "支持精确或包含查询，以 ERP 活动编号为准。"),
        ("查询条件", "活动时间", "按活动起止时间筛选；与筛选时间区间存在交集的活动应被返回。"),
        ("列表字段", "活动名称、活动编号、组合数、组合固定价、活动时间、适用门店、活动状态、最后修改时间", "字段来自 ERP 同步数据；组合固定价展示“起”价或代表价时，应明确展示口径。"),
        ("操作", "查询、重置、查看", "重置清空全部查询条件并恢复默认列表；查看进入只读详情。"),
    ], [1400, 3300, 4660])

    add_h2(doc, "5.3 管理后台 - 活动详情")
    add_body(doc, "详情页用于核对 ERP 已配置的活动规则和组合商品，不提供编辑、删除、启停或 ERP 配置回写。")
    add_table(doc, ["信息分组", "展示字段"], [
        ("基本信息", "活动名称、活动编号、活动状态、活动时间、适用门店、适用对象、参与日期/星期、最后修改时间。"),
        ("组合信息", "组合编号、组合名称、组合固定价、原价合计、限购规则、组合状态。"),
        ("组合商品明细", "商品编码、商品名称、规格、每组数量、销售价、商品状态；同一组合下的商品按 ERP 下发顺序展示。"),
        ("价格分摊信息", "如 ERP 下发每商品分摊金额或分摊规则，后台只读展示；该信息用于订单、退款和对账核验。"),
    ], [2500, 6860])

    add_h2(doc, "5.4 小程序商城 - 展示与凑单规则")
    add_table(doc, ["页面/场景", "展示与交互规则", "生效条件"], [
        ("商品列表/商品详情", "命中有效组合的商品展示“组合价 ¥X 起”标识；商品详情展示组合商品、每组数量、组合固定价及“加入组合/去凑单”入口。", "仅在当前门店、当前时间、适用对象和商品状态均满足时展示。"),
        ("组合详情/凑单", "清晰列出组合内每个商品、需购数量与已选数量；可一键加入尚缺商品，商品不可售时给出原因提示。", "同一组合内所有商品均可售且库存满足其数量时允许加入。"),
        ("购物车", "未凑齐时提示“再购 X 件可享组合价”；凑齐后以组合卡片展示组合固定价与节省金额。", "商品必须属于同一履约门店与同一有效组合；跨门店商品不合并凑单。"),
        ("结算页", "展示实际生效的组合、组合价及优惠金额；若资格在结算前失效，刷新为当前可用价格并提示变化原因。", "提交订单前必须再次校验活动、商品、数量、库存和用户适用资格。"),
    ], [1900, 4700, 2760])
    add_body(doc, "商城文案统一使用“组合价 ¥X”“已选 A/B 件”“再购 X 件享组合价”。组合未凑齐时不得将组合固定价作为单品成交价或订单优惠展示。")

    add_h2(doc, "5.5 价格生效、订单与售后规则")
    add_table(doc, ["规则项", "规则说明"], [
        ("完整组合计价", "同一订单中，指定商品及每组数量完整满足时，整组按 ERP 下发的组合固定价计价；不完整部分按其可用的其他价格规则计价。"),
        ("多组命中", "同一购物车可组成多组时，系统按可完整组成的最大组数计算；剩余商品不参与组合价。"),
        ("会员价比较", "同一完整组合同时具备组合价与会员价时，比较“组合固定价”与“该组合商品按会员价计算的合计金额”，整组按更低金额生效，不允许跨两种价格拼接。"),
        ("其他促销", "同一商品同一时间仅参加一个促销活动；组合价与优惠券、满减等权益是否可叠加，以 ERP 下发的叠加配置为准。未下发时默认不叠加。"),
        ("价格分摊", "订单须保存组合价到商品行的分摊金额或分摊规则，所有行分摊金额之和必须等于组合固定价；分摊与舍入规则由 ERP 统一下发并在商城、订单、退款端一致执行。"),
        ("库存与履约", "结算前按组合商品及数量校验库存和履约门店；任一商品库存不足或不支持当前履约方式时，该组合不生效。"),
        ("订单与退款", "订单保存活动编号、组合编号、命中组数、商品数量与分摊信息。部分退款、整组退款和换货的价格恢复/优惠回收口径由 ERP 订单规则确定，商城按该结果展示。"),
    ], [2500, 6860])

    add_h2(doc, "5.6 状态、边界与待确认项")
    for item in (
        "活动状态展示为未开始、进行中、已结束、已停用；状态由 ERP 启停状态与活动起止时间共同决定。",
        "ERP 配置变更后的同步时效、失败重试、字段映射和告警策略由技术方案确认；同步未完成时，商城使用最近一次成功同步数据，并在结算时再次校验。",
        "本期管理后台不包含组合价活动的新建、编辑、复制、删除、审批、权限配置、启停、提前结束及 ERP 数据回写。",
        "待确认：组合价与优惠券/满减的叠加范围、限购维度（用户/订单/活动）、分摊及尾差舍入规则、部分退款或换货后的优惠回收规则。确认后需同步更新 ERP、商城、订单和售后规则。",
    ):
        add_bullet(doc, item)

    add_h2(doc, "5.7 验收标准")
    add_table(doc, ["编号", "验收项", "通过标准"], [
        ("AC-01", "ERP 维护", "ERP 可维护活动、组合固定价、组合商品及数量、门店、时间和适用对象，并在启用前拦截不完整组合、商品冲突或不合法价格。"),
        ("AC-02", "数据同步", "后台与商城可读取 ERP 最近一次成功同步的活动、组合及商品明细；同步失败具备可追踪告警或错误记录。"),
        ("AC-03", "后台查询", "可按活动名称、活动编号和活动时间查询；重置后恢复默认列表；详情完整展示活动与组合商品且不可编辑。"),
        ("AC-04", "商城展示", "有效组合商品在商品详情、凑单页和购物车按规则展示组合价、需购数量和凑单进度；无效或不可售时不展示可成交的组合价。"),
        ("AC-05", "组合计价", "仅完整组合按组合固定价生效；购物车存在多组时按最大完整组数计算；未凑齐商品不得错误享受组合价。"),
        ("AC-06", "价格一致", "商品展示、购物车、结算页和订单中的组合价、会员价比较结果一致；整组会员价更低时整组按会员价生效。"),
        ("AC-07", "订单追溯", "订单可追溯活动编号、组合编号、命中组数和商品行分摊金额/规则，分摊合计等于组合固定价。"),
        ("AC-08", "边界控制", "管理后台无组合价维护入口；ERP 是活动及组合数据唯一维护来源，商城与后台不回写 ERP。"),
    ], [1100, 2500, 5760])

    doc.core_properties.title = "组合价需求文档"
    doc.core_properties.subject = "千金大药房组合价需求"
    doc.core_properties.author = "千金大药房"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_document()
