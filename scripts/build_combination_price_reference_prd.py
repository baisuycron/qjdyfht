from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
# The bundled Windows Python runtime can receive a mojibake absolute cwd when
# launched from a Chinese-named folder. Keep the build target relative and let
# PowerShell apply the final Chinese filename after generation.
OUTPUT = Path("combination_price_reference_v22.docx")

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
GRAY = "666666"
MUTED = "7A8491"
HEADER_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
CAUTION_FILL = "FFF7E3"
RISK_FILL = "FFF0F0"
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(run, size=11, bold=False, color="000000", italic=False):
    run.font.name = "Calibri"
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), "Calibri")
    rpr.rFonts.set(qn("w:hAnsi"), "Calibri")
    rpr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def configure_style(style, size, color="000000", bold=False, before=0, after=6, line=1.1):
    style.font.name = "Calibri"
    rpr = style._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), "Calibri")
    rpr.rFonts.set(qn("w:hAnsi"), "Calibri")
    rpr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)
    paragraph = style.paragraph_format
    paragraph.space_before = Pt(before)
    paragraph.space_after = Pt(after)
    paragraph.line_spacing = line


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    no_split = OxmlElement("w:cantSplit")
    tr_pr.append(no_split)


def set_table_geometry(table, widths, indent=TABLE_INDENT_DXA):
    if sum(widths) != CONTENT_WIDTH_DXA:
        raise ValueError(f"table widths must total {CONTENT_WIDTH_DXA}, got {sum(widths)}")
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    for tag, value in (("w:tblW", sum(widths)), ("w:tblInd", indent)):
        node = tbl_pr.first_child_found_in(tag)
        if node is None:
            node = OxmlElement(tag)
            tbl_pr.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    for index, width in enumerate(widths):
        table._tbl.tblGrid.gridCol_lst[index].set(qn("w:w"), str(width))
    for row in table.rows:
        prevent_row_split(row)
        for index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_table_borders(table, color="D9DEE7", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def write_cell(cell, text, bold=False, color="000000", size=9.4, align=WD_ALIGN_PARAGRAPH.LEFT):
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.18
    set_run_font(paragraph.add_run(str(text)), size=size, bold=bold, color=color)


def add_table(doc, headers, rows, widths, center_columns=None, header_fill=HEADER_FILL):
    center_columns = set(center_columns or [])
    table = doc.add_table(rows=1, cols=len(headers))
    for index, header in enumerate(headers):
        set_cell_shading(table.rows[0].cells[index], header_fill)
        write_cell(table.rows[0].cells[index], header, bold=True, color=INK, size=9.3, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_repeat_table_header(table.rows[0])
    for values in rows:
        if len(values) != len(headers):
            raise ValueError("table row length does not match header length")
        cells = table.add_row().cells
        for index, value in enumerate(values):
            align = WD_ALIGN_PARAGRAPH.CENTER if index in center_columns else WD_ALIGN_PARAGRAPH.LEFT
            write_cell(cells[index], value, align=align)
    set_table_geometry(table, widths)
    set_table_borders(table)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(2)
    return table


def add_body(doc, text, bold_lead=None):
    paragraph = doc.add_paragraph(style="Normal")
    if bold_lead and text.startswith(bold_lead):
        set_run_font(paragraph.add_run(bold_lead), bold=True, color=INK)
        set_run_font(paragraph.add_run(text[len(bold_lead):]))
    else:
        set_run_font(paragraph.add_run(text))
    return paragraph


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    size = {1: 16, 2: 13, 3: 12}[level]
    color = BLUE if level < 3 else DARK_BLUE
    set_run_font(paragraph.add_run(text), size=size, bold=True, color=color)
    return paragraph


def next_num_id(numbering):
    ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    return max(ids, default=0) + 1


def next_abstract_num_id(numbering):
    ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    return max(ids, default=0) + 1


def create_list_numbering(doc, ordered=False):
    numbering = doc.part.numbering_part.element
    abstract_id = next_abstract_num_id(numbering)
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), "decimal" if ordered else "bullet")
    level.append(fmt)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1." if ordered else "•")
    level.append(level_text)
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    level.append(justification)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "720")
    indent.set(qn("w:hanging"), "360")
    p_pr.append(indent)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)
    num_id = next_num_id(numbering)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_num_id = OxmlElement("w:abstractNumId")
    abstract_num_id.set(qn("w:val"), str(abstract_id))
    num.append(abstract_num_id)
    numbering.append(num)
    return num_id


def add_list_item(doc, text, num_id):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.167
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.extend((ilvl, num_id_node))
    p_pr.append(num_pr)
    set_run_font(paragraph.add_run(text))
    return paragraph


def add_callout(doc, title, text, fill=CALLOUT_FILL, title_color=DARK_BLUE):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.08)
    paragraph.paragraph_format.right_indent = Inches(0.08)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.2
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), fill)
    p_pr.append(shading)
    borders = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "4")
        border.set(qn("w:color"), "D5DCE6")
        borders.append(border)
    p_pr.append(borders)
    set_run_font(paragraph.add_run(f"{title}："), size=10.2, bold=True, color=title_color)
    set_run_font(paragraph.add_run(text), size=10.2, color="333333")
    return paragraph


def add_meta_row(doc, label, value):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.15
    set_run_font(paragraph.add_run(f"{label}："), size=10.3, bold=True, color=INK)
    set_run_font(paragraph.add_run(value), size=10.3, color="333333")


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, text, end))


def add_masthead_rule(doc):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(7)
    paragraph.paragraph_format.space_after = Pt(8)
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "10")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), BLUE)
    borders.append(bottom)
    p_pr.append(borders)


def new_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    configure_style(doc.styles["Normal"], 11, after=6, line=1.1)
    configure_style(doc.styles["Heading 1"], 16, BLUE, True, before=16, after=8, line=1.0)
    configure_style(doc.styles["Heading 2"], 13, BLUE, True, before=12, after=6, line=1.0)
    configure_style(doc.styles["Heading 3"], 12, DARK_BLUE, True, before=8, after=4, line=1.0)
    bullets = create_list_numbering(doc, ordered=False)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    set_run_font(header.add_run("千金大药房商城 | 产品需求文档"), size=9, color=GRAY)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_after = Pt(0)
    set_run_font(footer.add_run("组合价需求文档  ·  第 "), size=9, color=GRAY)
    add_page_field(footer)
    set_run_font(footer.add_run(" 页"), size=9, color=GRAY)

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(8)
    kicker.paragraph_format.space_after = Pt(3)
    set_run_font(kicker.add_run("PRODUCT REQUIREMENTS DOCUMENT"), size=9.5, bold=True, color=BLUE)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(4)
    title.paragraph_format.keep_with_next = True
    set_run_font(title.add_run("千金小程序需求-组合价"), size=23, bold=True, color=INK)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(14)
    set_run_font(subtitle.add_run("后台组合价 · 商城组合购 · 小程序交易链路"), size=12.5, color=GRAY)

    add_meta_row(doc, "所属模块", "营销中心 > 平台促销 > 组合价；商城商品详情 > 组合购")
    add_meta_row(doc, "版本", "V2.2")
    add_meta_row(doc, "日期", "2026-08-26")
    add_meta_row(doc, "目标读者", "产品、运营、交互、前端、后端、订单、售后与测试团队")
    add_meta_row(doc, "维护边界", "千金大药房后台统一维护组合价活动；千金大药房商城读取并使用同一套活动、商品、组合单价、数量、组合价格和状态数据，商城端不单独维护组合价活动")
    add_meta_row(doc, "文档状态", "数据关系修订版（含待确认问题）")
    add_masthead_rule(doc)
    return doc, bullets


def build():
    doc, bullets = new_document()

    add_heading(doc, "一、需求背景", 1)
    add_body(doc, "千金大药房后台项目已经具备“组合价”入口、活动列表、查询、新增、编辑、查看和手动结束等原型功能；千金大药房商城项目已经具备商品详情优惠组合、组合商品列表、加入购物车、立即购买、购物车整组展示、确认订单和商品组合单价展示等原型能力。")
    add_body(doc, "与此同时，旧版微信小程序采用“主商品固定、其他商品可选、直接进入组合购订单预览”的历史流程；当前目标则是由千金大药房后台统一维护 2-9 个组合商品、商品组合单价、每组数量和活动组合价格，商城读取同一活动后支持商品详情展示、组合商品列表、购物车和立即购买。因此本需求需要统一后台维护、商城展示、结算、订单和售后口径。")
    add_callout(doc, "核心原则", "千金大药房后台的组合价活动是唯一活动来源。商城读取并使用同一套组合价活动数据，不能单独新增、编辑或改写活动、商品、组合单价、数量、组合价格和状态；下单时再结合最终履约门店的售价、库存和商品可售情况进行检查。")

    add_heading(doc, "二、需求目标", 1)
    for text in (
        "让运营人员能够按活动和商品维度查询组合活动，并完成符合权限和状态约束的新增、编辑、查看及结束操作。",
        "让消费者在商品详情、组合商品列表和购物车中清楚理解整组商品、组合价格、套数和实际成交商品。",
        "保证商品详情、购物车、确认订单、提交订单、订单详情和售后使用同一活动版本、履约门店、商品组合单价和成交金额口径。",
        "明确后台与商城使用同一套组合价活动数据，并明确组合价与会员价、优惠券、折扣、满减及其他促销之间的互斥关系。",
        "把后台原型、商城原型和旧版小程序流程之间的差异保留为待确认项，避免将演示数据或旧流程直接当成正式规则。",
    ):
        add_list_item(doc, text, bullets)

    add_heading(doc, "三、业务流程", 1)
    add_body(doc, "组合购从千金大药房后台维护开始，商城读取同一活动，经商城交易服务支撑小程序展示、购物车、订单预览、提交、订单记录和整单售后。")
    add_table(
        doc,
        ["步骤", "处理方", "业务动作", "产出"],
        [
            ("1", "运营后台", "创建或编辑组合价活动，配置活动名称、活动时间、2-9 个组合商品、每个商品的组合单价及每组数量；活动组合价格自动汇总。", "可保存的组合价活动"),
            ("2", "后台服务", "检查活动名称和活动编号是否重复，并检查活动时间、商品数量、组合单价和促销占用；重复点击保存时不能重复创建活动。", "已发布的组合价活动"),
            ("3", "小程序", "从商城服务读取后台已发布的同一组合价活动，并按当前门店、商品和中国标准时间判断是否展示；无可用活动时不展示入口。", "组合购入口与活动列表"),
            ("4", "消费者", "在组合商品列表查看整组内容，选择加入购物车或立即购买；一次加入增加 1 套完整组合。", "整组购物车行或直接购买请求"),
            ("5", "订单预览", "根据后台同一活动和最终履约门店，重新检查实际商品清单、门店原价、活动组合价格、商品组合单价、库存和商品行成交金额。", "订单预览结果与本次预览凭证"),
            ("6", "订单提交", "再次检查活动、门店、实际商品清单和库存，并防止重复下单；检查通过后保存本次订单的成交信息。", "组合购订单"),
            ("7", "订单/售后", "展示组合标识、套数、后台商品组合单价和成交金额；含组合购订单只允许整单退款。", "可追溯订单与售后主单"),
            ("8", "生命周期", "后台活动自然结束、人工结束或因全局可参与商品不足 2 个而自动结束后，商城立即停止展示和接收新订单。", "一致的已结束状态与操作记录"),
        ],
        [720, 1500, 4620, 2520],
        center_columns={0, 1},
    )
    add_callout(doc, "冲突处理", "商城商品只要存在于任一未开始或进行中的组合购或其他促销中即判定冲突，不因活动时间不重叠而放行；已结束活动释放占用。组合购成交不与会员价、优惠券或其他营销活动叠加、比较或择优。", fill=CAUTION_FILL, title_color="7A5A00")

    add_heading(doc, "四、产品方案概要", 1)
    add_table(
        doc,
        ["范围", "方案说明"],
        [
            ("后台活动维护", "千金大药房后台提供活动查询、新增、编辑、查看和结束；每个商品配置组合单价和数量，活动组合价格由所有商品的“组合单价 × 数量”相加得到。"),
            ("商城活动使用", "商城读取后台已发布的同一活动，直接使用活动名称、活动时间、组合商品、组合单价、每组数量、组合价格和活动状态，不另建一套组合活动。"),
            ("同一数据来源", "后台活动新增、编辑、结束或商品内容变化后，商城刷新并使用最新活动数据；商城端不能单独修改活动内容。"),
            ("小程序端", "进行中的活动在商品详情展示“优惠组合”，进入独立组合商品列表，支持加入购物车和立即购买；整组商品、套数和价格不可拆分编辑。"),
            ("价格与结算", "商城直接使用后台活动中的组合价格、商品组合单价和每组数量；最终履约门店实时普通售价仅用于展示原价和计算优惠，原价变化不修改后台组合价格。"),
            ("订单与售后", "订单保存后台活动信息、履约门店、实际商品清单、门店原价、组合价格和商品行成交金额；含组合购订单仅支持整单退款。"),
        ],
        [1800, 7560],
    )

    add_heading(doc, "五、产品方案细节", 1)

    add_heading(doc, "5.1 组合价 - 功能入口", 2)
    add_body(doc, "后台入口路径：营销中心 > 平台促销 > 组合价。小程序入口位于参与活动商品的商品详情页“优惠组合”卡片，点击“去购买组合套装”进入独立的组合商品列表页面。")
    add_table(
        doc,
        ["入口", "展示条件", "交互结果"],
        [
            ("后台营销入口", "角色拥有组合价查看权限；新增、编辑和结束权限需按操作权限分别控制。", "进入组合价活动列表。"),
            ("商品详情优惠组合", "活动为进行中；当前商品属于活动；当前门店至少有 2 个可交易活动商品。", "展示组合名称、商品数量、后台活动组合价格和组合商品预览。"),
            ("去购买组合套装", "商品详情优惠组合可用。", "进入组合商品列表，不打开弹窗、不直接下单。"),
            ("入口异常", "活动未开始、已结束、当前门店可交易商品少于 2 个或服务不可用。", "隐藏组合购入口，不影响普通商品购买。"),
        ],
        [2100, 4440, 2820],
    )

    add_heading(doc, "5.2 组合价 - 活动列表", 2)
    add_body(doc, "页面默认展示后台当前可查询的组合价活动；所有查询条件同时生效，空值代表不限制，查询后回到第 1 页。")
    add_table(
        doc,
        ["区域", "字段/动作", "规则"],
        [
            ("查询条件", "活动时间", "按活动起止时间与筛选区间是否有交集进行过滤；不设默认值。"),
            ("查询条件", "活动名称", "全模糊查询，最多 100 个字符。"),
            ("查询条件", "活动状态", "下拉单选：全部、未开始、进行中、已结束。"),
            ("查询条件", "商品名称", "在组合商品中全模糊查询，最多 100 个字符。"),
            ("查询条件", "商品编码", "在组合商品中精确查询，最多 50 个字符。"),
            ("操作", "查询", "使用全部已填写条件触发查询。"),
            ("操作", "重置", "清空全部条件和页码，恢复默认列表。"),
            ("列表", "展示字段", "活动名称、活动时间、活动状态、操作；不在列表展示商品明细、营销类型、规格/单位、厂家或组合商品数量。"),
            ("分页", "页码", "支持 10、20、50、100 条/页、上一页、下一页和页码跳转。"),
        ],
        [1500, 2100, 5760],
    )
    add_table(
        doc,
        ["活动状态", "可用操作", "说明"],
        [
            ("未开始", "查看、编辑、结束", "编辑复用新增表单并回填全部字段。"),
            ("进行中", "查看、结束", "不得编辑活动配置。"),
            ("已结束", "查看", "包括自然结束、人工结束和系统自动结束。"),
        ],
        [1800, 2400, 5160],
        center_columns={0, 1},
    )

    add_heading(doc, "5.3 组合价 - 活动列表 - 活动详情", 2)
    add_body(doc, "后台维护的组合价活动就是商城使用的组合购活动，两端使用同一活动编号和同一套活动数据。后台负责维护，商城负责展示和交易，商城端没有单独的组合价活动配置入口。")
    add_table(
        doc,
        ["环节", "配置/展示字段", "规则"],
        [
            ("后台活动信息", "活动名称、活动编号、开始时间、结束时间、只读组合价格、组合商品。", "活动编号由系统生成；活动组合价格由所有商品的“组合单价 × 数量”相加得到。组合描述字段已取消。"),
            ("后台商品明细", "商品图片、商品编码、商品名称、最新售价、组合单价、商品规格、数量、操作。", "2-9 个商品；组合单价大于 0 且不高于最新售价；数量为正整数；新增商品默认组合单价等于最新售价。"),
            ("商城活动展示", "活动名称、活动编号、活动时间、活动状态、组合价格、组合商品、商品组合单价及每组数量。", "全部取自后台同一组合价活动；商城只能读取和展示，不能新增、编辑、替换或删除活动商品。"),
            ("后台详情页", "复用新增和编辑页的字段结构。", "只读展示；无商品选择、数量编辑、删除和保存。人工结束需保留计划结束时间并记录操作人、操作时间。"),
        ],
        [1800, 3300, 4260],
    )
    add_table(
        doc,
        ["校验项", "后台维护规则", "商城使用规则"],
        [
            ("时间", "结束时间晚于开始时间且晚于保存时当前时间；开始时间早于当前时，保存后立即进入进行中。", "直接使用后台活动时间和状态，未开始不展示，进行中才可成交，已结束停止展示。"),
            ("商品数量", "至少 2 个、最多 9 个，不区分主商品。", "使用后台活动中的同一商品清单和每组数量，消费者不能任选、替换或删除单个商品。"),
            ("价格", "每个商品的组合单价大于 0 且不高于最新售价；活动组合价格自动汇总。", "直接使用后台活动中的组合价格、商品组合单价和数量，商城端不得重新配置或改写。"),
            ("促销冲突", "保存时检查未开始和进行中的组合价及其他促销活动，避免同一商品重复占用。", "只使用后台检查通过并已发布的活动；发现活动失效或商品不可交易时停止本次预览或下单。"),
            ("保存与生效", "重复点击保存不能重复创建活动；活动、商品、价格、活动版本和操作记录要么全部成功，要么全部失败。", "后台保存成功后，商城读取该活动的最新版本；读取失败时不展示旧活动价格。"),
        ],
        [1500, 3810, 4050],
    )

    add_heading(doc, "5.4 小程序端 - 展示规则", 2)
    add_table(
        doc,
        ["页面/区域", "展示内容", "交互与规则"],
        [
            ("商品详情", "“组合购”标识、“优惠组合”卡片、组合商品预览、组合价格。", "进行中且当前门店可用时展示；活动、商品、组合价格和数量均取自后台同一组合价活动。"),
            ("组合商品列表", "活动名称、商品数量、组合价格、商品图片/名称/规格、商品组合单价、每组数量。", "支持加入购物车和立即购买；消费者不能勾选、删除、替换单个组合商品。"),
            ("加入购物车", "以一个组合行展示活动，内部商品可展开查看。", "每点击一次增加 1 套完整组合；同活动版本、同履约门店合并套数；角标按组合套数增加。"),
            ("购物车整组", "组合价、套数、组合商品总件数、后台商品组合单价和商品行成交金额。", "整组勾选、删除和修改套数；不能只处理其中部分商品。增加套数前按增加后的总套数重新校验。"),
            ("门店变化", "重新展示该门店可交易商品、原价与可用状态。", "地址、定位、配送方式或门店变化使旧预览失效；最终履约门店少于 2 个可交易商品时隐藏入口。"),
            ("原价不高于组合价", "仅展示后台活动组合价格。", "不展示节省金额、优惠金额或原价划线，但活动继续生效。"),
        ],
        [1800, 3300, 4260],
    )
    add_callout(doc, "后台与商城数据关系", "商城组合购来自后台组合价活动。活动编号、名称、时间、状态、组合商品、商品组合单价、每组数量和组合价格必须保持一致；商城端只负责展示和交易，不能另建活动或修改后台配置。")

    add_heading(doc, "5.5 交易订单相关", 2)
    add_table(
        doc,
        ["环节", "规则说明"],
        [
            ("立即购买", "从组合商品列表进入确认订单，并带上后台组合价活动编号、活动版本、履约门店和购买套数。"),
            ("购物车结算", "整组生成订单商品行；若最终允许与普通商品混合结算，组合部分仍保持活动和售后边界。混单规则见 5.9 待确认。"),
            ("订单预览", "系统读取后台同一组合价活动，重新检查实际商品清单、最终履约门店普通售价、活动组合价格、商品组合单价、每组数量和库存，并返回本次订单预览凭证。"),
            ("订单提交", "使用本次订单预览凭证，并防止用户重复点击产生多个订单；再次检查后台活动是否有效、活动版本是否变化、门店和商品清单是否变化、库存是否足够。任何一项变化都要重新预览并由用户确认。"),
            ("订单记录", "保存后台活动编号、活动名称、活动版本、履约门店、配置商品和实际成交商品、商品组合单价、每组数量、购买套数、门店原价、组合价格和商品行成交金额。"),
            ("正向订单", "列表和详情展示“组合购”标识；详情展示每套组合价、购买套数、实际商品与商品行成交金额。"),
            ("退款", "只要订单含组合购，退款范围固定为整笔订单全部商品和全部数量；不支持部分商品、部分数量或部分套数退款。"),
            ("异常售后", "缺货、漏发、破损、召回和配送丢失通过补发、客服补偿或整单售后主单处理；跨仓/包裹子任务不得独立退款。"),
        ],
        [2100, 7260],
    )

    add_heading(doc, "5.6 报表相关", 2)
    add_table(
        doc,
        ["指标/报表", "统计口径"],
        [
            ("活动效果", "活动曝光、详情点击、加入购物车套数、立即购买次数、预览成功率、下单转化率、订单数、组合套数和组合购成交额。"),
            ("价格优惠", "以订单记录中的门店原价合计与后台活动组合成交总额计算；原价合计不高于组合价时优惠金额记 0，不生成负优惠。"),
            ("商品成交金额", "报表按后台商品组合单价、每组数量和购买套数统计商品成交金额；所有商品行合计必须等于组合购成交总额。"),
            ("活动生命周期", "使用后台活动的自然结束、人工结束和系统自动结束结果分别统计；记录结束原因、操作人或触发原因和时间。"),
            ("门店可用性", "统计因单店库存不足隐藏入口、库存恢复重新展示、全局资格商品不足 2 个导致自动结束的次数。"),
            ("售后", "整单退款率、退款金额、异常补发/补偿、跨仓售后时长和退款责任类型。"),
            ("对账", "订单、支付、发票、财务和退款报表均使用同一订单成交快照，不按活动当前配置回算。"),
        ],
        [2400, 6960],
    )

    add_heading(doc, "5.7 价格生效规则", 2)
    add_table(
        doc,
        ["规则项", "规则说明"],
        [
            ("组合价来源", "商城直接使用后台同一活动中保存的组合价格、商品组合单价和每组数量，商城端不能另设一套价格。"),
            ("组合成交总额", "每套都按后台活动中的组合价格成交。买几套，就用“每套组合价 × 购买套数”计算总金额。"),
            ("原价来源", "以最终履约门店中实际成交商品的实时普通售价计算，不使用会员价或其他促销价。"),
            ("原价合计", "每个商品都按“门店售价 × 每组数量 × 购买套数”计算，再把所有商品金额加起来。商品原价变化不会修改后台活动中的组合价，也不会让活动自动结束。"),
            ("商品成交金额", "每个商品行按“后台配置的商品组合单价 × 每组数量 × 购买套数”计算。所有商品行相加必须等于组合成交总金额。"),
            ("金额一致", "后台保存活动前先检查商品行金额合计与活动组合价格是否一致；商城预览和提交订单时再次检查，不一致时停止下单并刷新后台活动数据。"),
            ("营销互斥", "会员价、优惠券、折扣、满减、积分抵现及其他营销不参与组合购展示、比较、锁定、核销或结算。"),
            ("展示与结算一致", "商品详情只做参考展示；进入购物车、确认订单和提交订单时都要重新检查，最终以系统返回的成交结果为准。"),
            ("价格变化提示", "后台活动版本、组合价格、商品清单、商品组合单价、门店或售价发生变化时，原来的订单预览结果失效；系统读取最新后台活动并提示用户重新确认。"),
        ],
        [2100, 7260],
    )
    add_callout(doc, "计算示例", "后台活动中 A、B、C 三个商品每组各 1 件，商品组合单价分别为 48 元、24 元和 8 元，活动组合价格自动汇总为 80 元。购买 2 套时，三个商品行成交金额分别为 96 元、48 元和 16 元，合计 160 元，与“80 元 × 2 套”一致。")

    add_heading(doc, "5.8 状态与边界", 2)
    add_table(
        doc,
        ["边界项", "规则"],
        [
            ("自然状态", "按中国标准时间计算：当前时间早于开始时间为未开始；开始时间至计划结束时间为进行中；晚于计划结束时间为已结束。"),
            ("状态覆盖", "只要活动已被人工结束或系统自动结束，就优先显示“已结束”；原计划结束时间仍保留，不被改写。"),
            ("人工结束", "后台中的未开始和进行中活动均可结束；确认后商城立即隐藏入口并拒绝新预览和下单，记录操作人和时间。结束前已创建的待支付订单按订单自身超时规则处理。"),
            ("单店不足", "仅某门店因库存不足导致可交易商品少于 2 个时，只隐藏该门店入口；库存恢复且活动仍有效时可恢复。"),
            ("全局不足", "因下架、停售、删除或全局失效导致具备活动资格的商品少于 2 个时，后台活动自动结束、释放促销占用且不自动恢复；商城随即停止展示。"),
            ("本期不含", "活动复制、删除、恢复、审批、批量导入、指定门店/会员、消费者任选/替换商品、部分退款和跨订单凑单。"),
            ("实现状态", "后台和商城当前页面仍是前端演示，旧版微信小程序使用的是历史购买流程。正式上线前需补齐“后台统一维护、商城读取同一活动”的活动查询、价格、库存、订单和售后服务。"),
        ],
        [2100, 7260],
    )

    add_heading(doc, "5.9 待确认问题", 2)
    add_table(
        doc,
        ["编号", "待确认问题", "影响"],
        [
            ("P-01", "含组合购的购物车是否允许同时结算普通商品或另一组合活动？", "当前商城原型可以出现混合商品，最新商城口径倾向一个活动单独下单。"),
            ("P-02", "旧版微信小程序的主商品固定、其他商品可选模式是否整体废弃？", "决定旧流程是否下线，以及用户选择商品和直接进入订单预览的方式是否保留。"),
            ("P-03", "每组商品数量在下单时按什么方式传递？", "需要保证购买多套时，每个商品数量都按“每组数量 × 购买套数”准确计算。"),
            ("P-04", "某门店的部分组合商品因库存、下架、停售、禁限售、处方或配送规则不能交易，但剩余商品仍不少于 2 个时，是去掉不可交易商品继续购买，还是整组不允许购买？如继续购买，活动组合价格及被去掉商品的组合单价怎么处理？", "直接影响商城展示、商品行成交金额、活动组合价格、订单和对账是否一致。"),
            ("P-05", "后台正式保存活动、生成活动编号、记录操作和自动结束活动，以及商城读取同一活动，分别由哪个服务提供？", "当前页面只是前端演示，需要明确后台保存与商城读取同一活动的正式接口。"),
        ],
        [1050, 4650, 3660],
        center_columns={0},
        header_fill=CAUTION_FILL,
    )
    add_callout(doc, "评审要求", "P-01、P-02、P-03 和 P-04 会直接改变商城购物车、订单、商品数量或成交金额计算方式，应在研发排期前确认；在确认前不得把当前演示数据、旧版小程序流程或后台原型字段描述成已确定的正式规则。", fill=RISK_FILL, title_color="9B1C1C")

    add_heading(doc, "验收基线", 3)
    add_table(
        doc,
        ["编号", "验收项", "通过标准"],
        [
            ("AC-01", "同一活动来源", "组合价活动由千金大药房后台统一维护；商城读取并使用同一活动编号、商品、组合单价、数量、组合价格和状态，商城端不能单独修改。"),
            ("AC-02", "入口", "仅进行中且当前门店至少 2 个可交易商品时展示组合购入口；异常不阻断普通购买。"),
            ("AC-03", "整组购物车", "每次加入增加 1 套，同一个活动、同一个门店的组合合并购买套数，内部商品不可拆分操作。"),
            ("AC-04", "门店重算", "地址、定位或门店变化后重新查询后台活动、商品、售价和库存，并重新计算商品行成交金额，原订单预览结果失效。"),
            ("AC-05", "价格互斥", "组合购不与会员价、优惠券或其他营销叠加、比较或择优。"),
            ("AC-06", "商品成交金额", "每个商品行按后台商品组合单价、每组数量和购买套数计算，所有商品行合计严格等于组合成交总额。"),
            ("AC-07", "订单记录", "活动结束或商品删除后，历史订单仍完整展示后台活动编号、成交商品、门店、商品组合单价和组合价格。"),
            ("AC-08", "状态", "后台活动人工结束或自动结束后，商城立即停止展示并阻止新订单，同时保留计划结束时间并记录原因和操作。"),
            ("AC-09", "整单退款", "含组合购订单无法选择部分商品、数量或套数退款。"),
            ("AC-10", "系统重新计算", "修改前端传来的活动编号、门店、商品、数量或价格，不能改变系统根据后台同一活动重新计算后的成交结果。"),
        ],
        [1050, 2100, 6210],
        center_columns={0},
    )

    doc.core_properties.title = "千金小程序需求-组合价"
    doc.core_properties.subject = "后台组合价与商城组合购需求"
    doc.core_properties.author = "千金大药房产品团队"
    doc.core_properties.keywords = "组合价, 组合购, 后台, 商城, 同一活动, 组合单价, 整单退款, PRD"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
